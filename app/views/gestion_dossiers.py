#!/usr/bin/env python
#-*- coding: utf-8

# Imports
from app.decorators import *
from app.forms.gestion_dossiers import GererFicheVie
from app.functions import ger_droits
from app.functions import init_fm
from app.functions import rempl_fich_log
from app.models import TFicheVie

from django.views.decorators.csrf import csrf_exempt
from django.core.urlresolvers import reverse
from django.http import HttpResponse
from django.shortcuts import get_object_or_404
from django.shortcuts import render

import json

'''
Cette vue permet d'afficher le menu principal du module de gestion des dossiers.
request : Objet requête
'''
@verif_acc
def index(request) :

	# Imports
	from app.functions import get_thumbnails_menu
	from django.http import HttpResponse
	from django.shortcuts import render

	output = HttpResponse()

	if request.method == 'GET' :

		# J'affiche le template.
		output = render(request, './gestion_dossiers/main.html', {
			'menu' : get_thumbnails_menu('gest_doss', 2), 'title' : 'Gestion des dossiers'
		})

	return output

'''
Cette vue permet d'afficher la page de création d'un dossier ou de traiter une action.
request : Objet requête
'''
@verif_acc
def cr_doss(request) :

	# Imports
	from app.forms.gestion_dossiers import GererDossier
	from app.functions import alim_ld
	from app.functions import dt_fr
	from app.functions import filtr_doss
	from app.functions import gen_t_ch_doss
	from app.functions import init_f
	from app.functions import init_fm
	from app.functions import rempl_fich_log
	from app.models import TAction
	from app.models import TAxe
	from app.models import TDossier
	from app.models import TFamille
	from app.models import TMoa
	from app.models import TProgramme
	from app.models import TSousAxe
	from app.models import TTypesProgrammesTypeDossier
	from django.core.urlresolvers import reverse
	from django.http import HttpResponse
	from django.shortcuts import render
	import json

	output = HttpResponse()

	if request.method == 'GET' :

		# J'instancie un objet "formulaire".
		f_cr_doss = GererDossier(prefix = 'GererDossier', k_util = request.user)

		# Je déclare un tableau qui stocke le contenu de certaines fenêtres modales.
		t_cont_fm = {
			'ch_doss' : gen_t_ch_doss(request)
		}

		# Je déclare un tableau de fenêtres modales.
		t_fm = [
			init_fm('ch_doss', 'Choisir un dossier associé', t_cont_fm['ch_doss']),
			init_fm('cr_doss', 'Créer un dossier')
		]

		# J'affiche le template.
		output = render(
			request,
			'./gestion_dossiers/cr_doss.html',
			{ 'f_cr_doss' : init_f(f_cr_doss), 't_fm' : t_fm, 'title' : 'Créer un dossier' }
		)

	else :
		if 'action' in request.GET :

			# Je stocke la valeur du paramètre "GET" dont la clé est "action".
			get_action = request.GET['action']

			# Je traite le cas où je dois alimenter les listes déroulantes des axes, des sous-axes, des actions et des
			# types de dossiers.
			if get_action == 'alimenter-listes' :

				# J'affiche ou je cache certaines listes déroulantes selon les données qu'elles contiennent.
				output = HttpResponse(json.dumps(alim_ld(request)), content_type = 'application/json')

			# Je traite le cas où je dois filtrer les dossiers selon certains critères.
			if get_action == 'filtrer-dossiers' :

				# Je prépare le tableau des dossiers filtrés.
				t_doss = [(
					d.num_doss,
					d.get_int_doss(),
					d.id_org_moa.n_org,
					dt_fr(d.dt_delib_moa_doss) or '-',
					'<span class="choose-icon pointer pull-right" title="Choisir le dossier"></span>'
				) for d in filtr_doss(request)]

				# J'envoie le tableau des dossiers filtrés.
				output = HttpResponse(
					json.dumps({ 'success' : { 'datatable' : t_doss }}), content_type = 'application/json'
				)

		else :

			# Je soumets le formulaire.
			f_cr_doss = GererDossier(request.POST, request.FILES, prefix = 'GererDossier', k_util = request.user)

			# Je rajoute des choix valides pour certaines listes déroulantes (prévention d'erreurs).
			post_progr = request.POST['GererDossier-zl_progr']
			post_axe = request.POST['GererDossier-zl_axe']
			post_ss_axe = request.POST['GererDossier-zl_ss_axe']
			if post_progr :
				t_axes = [(a.pk, a.pk) for a in TAxe.objects.filter(id_progr = post_progr)]
				f_cr_doss.fields['zl_axe'].choices = t_axes
				t_types_doss = [
					(td.id_type_doss.pk, td.id_type_doss.pk) for td in TTypesProgrammesTypeDossier.objects.filter(
						id_type_progr = TProgramme.objects.get(pk = post_progr).id_type_progr.pk
					)
				]
				f_cr_doss.fields['zl_type_doss'].choices = t_types_doss
				t_ss_axes = [(sa.pk, sa.pk) for sa in TSousAxe.objects.filter(id_axe = post_axe)]
				f_cr_doss.fields['zl_ss_axe'].choices = t_ss_axes
				t_act = [(a.pk, a.pk) for a in TAction.objects.filter(id_ss_axe = post_ss_axe)]
				f_cr_doss.fields['zl_act'].choices = t_act

			if f_cr_doss.is_valid() :

				# Je récupère les données du formulaire valide.
				cleaned_data = f_cr_doss.cleaned_data
				v_doss_ass = cleaned_data.get('za_doss_ass')
				v_org_moa = cleaned_data.get('zl_org_moa')
				v_progr = cleaned_data.get('zl_progr')

				# Je pointe vers l'objet TMoa.
				o_org_moa = TMoa.objects.get(pk = v_org_moa)
				o_progr = TProgramme.objects.get(pk = v_progr)

				# J'initialise la valeur de la famille (valeur existante ou nouvelle instance).
				if v_doss_ass :
					o_fam = TDossier.objects.get(num_doss = v_doss_ass).id_fam
				else :
					o_fam = TFamille()
					o_fam.save()

				# Je prépare la valeur de chaque constituant du numéro de dossier.
				dim_progr = o_progr.dim_progr
				if o_org_moa.dim_org_moa :
					dim_org_moa = o_org_moa.dim_org_moa
				else :
					dim_org_moa = 'X'
				seq_progr = str(o_progr.seq_progr).zfill(2)

				# Je créé la nouvelle instance TDossier.
				o_nv_doss = f_cr_doss.save(commit = False)
				o_nv_doss.num_doss = '{0}-{1}-{2}'.format(dim_progr, dim_org_moa, seq_progr)
				o_nv_doss.id_fam = o_fam
				o_nv_doss.save()

				# Je mets à jour le séquentiel.
				o_progr.seq_progr = int(o_progr.seq_progr) + 1
				o_progr.save()

				# J'affiche le message de succès.
				output = HttpResponse(
					json.dumps({ 'success' : {
						'message' : 'Le dossier {0} a été créé avec succès.'.format(o_nv_doss),
						'redirect' : reverse('cons_doss', args = [o_nv_doss.id_doss])
					}}),
					content_type = 'application/json'
				)

				# Je complète le fichier log.
				rempl_fich_log([request.user.pk, request.user, o_nv_doss.pk, o_nv_doss, 'C', 'Dossier', o_nv_doss.pk])

			else :

				# J'affiche les erreurs.
				t_err = {}
				for k, v in f_cr_doss.errors.items() :
					t_err['GererDossier-{0}'.format(k)] = v
				output = HttpResponse(json.dumps(t_err), content_type = 'application/json')

	return output

'''
Cette vue permet d'afficher la page de modification d'un dossier ou de traiter le formulaire de mise à jour d'un
dossier.
request : Objet requête
_d : Identifiant d'un dossier
'''
@verif_acc
def modif_doss(request, _d) :

	# Imports
	from app.forms.gestion_dossiers import GererDossier
	from app.functions import alim_ld
	from app.functions import dt_fr
	from app.functions import filtr_doss
	from app.functions import gen_t_ch_doss
	from app.functions import ger_droits
	from app.functions import init_f
	from app.functions import init_fm
	from app.functions import rempl_fich_log
	from app.models import TAction
	from app.models import TAxe
	from app.models import TDossier
	from app.models import TDossierGeom
	from app.models import TFamille
	from app.models import TFinancement
	from app.models import TPrestationsDossier
	from app.models import TSousAxe
	from django.contrib.gis import geos
	from django.core.urlresolvers import reverse
	from django.http import HttpResponse
	from django.shortcuts import get_object_or_404
	from django.shortcuts import render
	import json

	output = HttpResponse()

	# Je vérifie l'existence d'un objet TDossier.
	o_doss = get_object_or_404(TDossier, pk = _d)

	# Je vérifie le droit d'écriture.
	ger_droits(request.user, o_doss, False)

	if request.method == 'GET' :

		# J'instancie un objet "formulaire".
		f_modif_doss = GererDossier(instance = o_doss, prefix = 'GererDossier')

		# Je déclare un tableau qui stocke le contenu de certaines fenêtres modales.
		t_cont_fm = {
			'ch_doss' : gen_t_ch_doss(request, o_doss.pk)
		}

		# Je déclare un tableau de fenêtres modales.
		t_fm = [
			init_fm('ch_doss', 'Choisir un dossier associé', t_cont_fm['ch_doss']),
			init_fm('modif_doss', 'Modifier un dossier')
		]

		# J'affiche le template.
		output = render(
			request,
			'./gestion_dossiers/modif_doss.html',
			{
				'd' : o_doss,
				'f_modif_doss' : init_f(f_modif_doss),
				't_fm' : t_fm,
				't_fin_length' : len(TFinancement.objects.filter(id_doss = o_doss)),
				't_prest_doss_length' : len(TPrestationsDossier.objects.filter(id_doss = o_doss)),
				'title' : 'Modifier un dossier'
			}
		)

	else :
		if 'action' in request.GET :

			# Je stocke la valeur du paramètre "GET" dont la clé est "action".
			get_action = request.GET['action']

			# Je traite le cas où je dois alimenter les listes déroulantes des axes, des sous-axes, des actions et des
			# types de dossiers.
			if get_action == 'alimenter-listes' :

				# J'affiche ou je cache certaines listes déroulantes selon les données qu'elles contiennent.
				output = HttpResponse(json.dumps(alim_ld(request)), content_type = 'application/json')

			# Je traite le cas où je dois filtrer les dossiers selon certains critères.
			if get_action == 'filtrer-dossiers' :

				# Je prépare le tableau des dossiers filtrés.
				t_doss = [(
					d.num_doss,
					d.get_int_doss(),
					d.id_org_moa.n_org,
					dt_fr(d.dt_delib_moa_doss) or '-',
					'<span class="choose-icon pointer pull-right" title="Choisir le dossier"></span>'
				) for d in filtr_doss(request, o_doss.pk)]

				# J'envoie le tableau des dossiers filtrés.
				output = HttpResponse(
					json.dumps({ 'success' : { 'datatable' : t_doss }}), content_type = 'application/json'
				)

			# Je traite le cas où je dois modifier la géométrie de l'objet TDossier.
			if get_action == 'modifier-geom' :

				# Je supprime les géométries existantes du dossier.
				TDossierGeom.objects.filter(id_doss = o_doss).delete()

				if request.POST['edit-geom'] :

					# Je récupère les objets créés.
					editgeom = request.POST['edit-geom'].split(';')

					# Je créé la nouvelle instance TDossierGeom.
					for g in editgeom :
						geom = geos.GEOSGeometry(g)
						geom_doss = TDossierGeom(id_doss = o_doss)
						if geom and isinstance(geom, geos.Polygon):
							geom_doss.geom_pol = geom
						if geom and isinstance(geom, geos.LineString):
							geom_doss.geom_lin = geom
						if geom and isinstance(geom, geos.Point):
							geom_doss.geom_pct = geom
						geom_doss.save()

				# J'affiche le message de succès.
				output = HttpResponse(
					json.dumps({ 'success' : {
						'message' : 'La géométrie du dossier {0} a été mise à jour avec succès.'.format(o_doss),
						'redirect' : reverse('cons_doss', args = [o_doss.pk])
					}}),
					content_type = 'application/json'
				)

				# Je complète le fichier log.
				rempl_fich_log([
					request.user.pk, request.user, o_doss.pk, o_doss, 'U', 'Géométrie d\'un dossier', o_doss.pk
				])

		else :

			# Je soumets le formulaire.
			f_modif_doss = GererDossier(request.POST, request.FILES, instance = o_doss, prefix = 'GererDossier')

			# Je rajoute des choix valides pour certaines listes déroulantes (prévention d'erreurs).
			post_progr = request.POST['GererDossier-zl_progr']
			post_axe = request.POST['GererDossier-zl_axe']
			post_ss_axe = request.POST['GererDossier-zl_ss_axe']
			if post_progr :
				t_axes = [(a.pk, a.pk) for a in TAxe.objects.filter(id_progr = post_progr)]
				f_modif_doss.fields['zl_axe'].choices = t_axes
				f_modif_doss.fields['zl_type_doss'].choices = [(o_doss.id_type_doss.pk, o_doss.id_type_doss)]
				t_ss_axes = [(sa.pk, sa.pk) for sa in TSousAxe.objects.filter(id_axe = post_axe)]
				f_modif_doss.fields['zl_ss_axe'].choices = t_ss_axes
				t_act = [(a.pk, a.pk) for a in TAction.objects.filter(id_ss_axe = post_ss_axe)]
				f_modif_doss.fields['zl_act'].choices = t_act

			if f_modif_doss.is_valid() :

				# Je récupère les données du formulaire valide.
				cleaned_data = f_modif_doss.cleaned_data
				v_doss_ass = cleaned_data.get('za_doss_ass')

				# Je pointe vers l'objet TDossier relatif au dossier associé.
				o_doss_ass = TDossier.objects.get(pk = o_doss.pk).id_doss_ass

				# J'initialise la valeur de la famille (valeur existante ou nouvelle instance).
				if v_doss_ass :
					v_fam = TDossier.objects.get(num_doss = v_doss_ass).id_fam
					TDossier.objects.filter(id_doss_ass = o_doss.pk).update(id_fam = v_fam)
				else :
					if o_doss_ass :
						v_fam = TFamille()
						v_fam.save()
					else :
						v_fam = o_doss.id_fam

				# Je modifie l'instance TDossier.
				o_doss_modif = f_modif_doss.save(commit = False)
				o_doss_modif.id_fam = v_fam
				o_doss_modif.save()

				# J'affiche le message de succès.
				output = HttpResponse(
					json.dumps({ 'success' : {
						'message' : 'Le dossier {0} a été mis à jour avec succès.'.format(o_doss_modif),
						'redirect' : reverse('cons_doss', args = [o_doss_modif.pk])
					}}),
					content_type = 'application/json'
				)

				# Je complète le fichier log.
				rempl_fich_log([
					request.user.pk, request.user, o_doss_modif.pk, o_doss_modif, 'U', 'Dossier', o_doss_modif.pk
				])

			else :

				# J'affiche les erreurs.
				t_err = {}
				for k, v in f_modif_doss.errors.items() :
					t_err['GererDossier-{0}'.format(k)] = v
				output = HttpResponse(json.dumps(t_err), content_type = 'application/json')

	return output

'''
Cette vue permet de traiter le formulaire de suppression d'un dossier.
request : Objet requête
_d : Identifiant d'un dossier
'''
@verif_acc
@csrf_exempt
def suppr_doss(request, _d) :

	# Imports
	from app.functions import ger_droits
	from app.functions import rempl_fich_log
	from app.models import TArretesDossier
	from app.models import TDossier
	from app.models import TDossierGeom
	from app.models import TDossierPgre
	from app.models import TFicheVie
	from app.models import TFinancement
	from app.models import TPhoto
	from app.models import TPrestationsDossier
	from django.core.urlresolvers import reverse
	from django.http import HttpResponse
	from django.shortcuts import get_object_or_404
	import json

	output = HttpResponse()

	# Je vérifie l'existence d'un objet TDossier.
	o_doss = get_object_or_404(TDossier, pk = _d)

	if request.method == 'POST' :

		# Je vérifie le droit d'écriture.
		ger_droits(request.user, o_doss, False)

		# J'initialise un tableau de jeu de données.
		t_qs = [
			['Événement', TFicheVie.objects.filter(id_doss = o_doss)],
			['Financement', TFinancement.objects.filter(id_doss = o_doss)],
			['Prestation', TPrestationsDossier.objects.filter(id_doss = o_doss)],
			['Arrêté', TArretesDossier.objects.filter(id_doss = o_doss)],
			['Photo', TPhoto.objects.filter(id_doss = o_doss)],
			['Objet géométrique', TDossierGeom.objects.filter(id_doss = o_doss)]
		]

		# Je vérifie si je peux exécuter ou non la suppression du dossier.
		peut_suppr = True
		t_elem_a_suppr = []
		for i in range(0, len(t_qs)) :
			if len(t_qs[i][1]) > 0 :
				peut_suppr = False
				cle = t_qs[i][0]
				if len(t_qs[i][1]) > 1 :
					split = cle.split(' ')
					for j in range(0, len(split)) :
						split[j] += 's'
					cle = ' '.join(split)
				t_elem_a_suppr.append('{0} : {1}'.format(cle, len(t_qs[i][1])))
		qs_act_pgre = TDossierPgre.objects.filter(id_doss = o_doss)
		if len(qs_act_pgre) > 0 :
			peut_suppr = False
			cle = 'Action PGRE'
			if len(qs_act_pgre) > 1 :
				cle = 'Actions PGRE'
			t_elem_a_suppr.append('{0} : {1}'.format(cle, len(qs_act_pgre)))

		if peut_suppr == True :

			# Je pointe vers l'objet TDossier à supprimer.
			o_doss_suppr = o_doss

			# Je récupère l'identifiant du dossier afin de le renseigner dans le fichier log.
			v_id_doss = o_doss.pk

			# Je supprime l'objet TDossier.
			o_doss.delete()

			# J'affiche le message de succès.
			output = HttpResponse(
				json.dumps({ 'success' : {
					'message' : 'Le dossier {0} a été supprimé avec succès.'.format(o_doss_suppr),
					'redirect' : reverse('ch_doss')
				}}),
				content_type = 'application/json'
			)

			# Je complète le fichier log.
			rempl_fich_log([request.user.pk, request.user, v_id_doss, o_doss_suppr, 'D', 'Dossier', v_id_doss])

		else :

			# Je prépare le message d'alerte.
			mess_html = 'Veuillez d\'abord modifier ou supprimer le/les élément(s) suivant(s) :'
			mess_html += '<ul class="list-inline">'
			for i in range(0, len(t_elem_a_suppr)) :
				mess_html += '<li>{0}</li>'.format(t_elem_a_suppr[i])
			mess_html += '</ul>'

			# J'affiche le message d'alerte.
			output = HttpResponse(json.dumps([mess_html]), content_type = 'application/json')

	return output

'''
Cette vue permet d'afficher la page de choix d'un dossier ou de traiter une action.
request : Objet requête
'''
@verif_acc
def ch_doss(request) :

	# Imports
	from app.forms.gestion_dossiers import ChoisirDossier
	from app.functions import alim_ld
	from app.functions import dt_fr
	from app.functions import filtr_doss
	from app.functions import init_f
	from app.functions import obt_doss_regr
	from app.models import TDossier
	from app.models import TMoa
	from app.models import TUtilisateur
	from django.core.urlresolvers import reverse
	from django.http import HttpResponse
	from django.shortcuts import render
	import json

	output = HttpResponse()

	if request.method == 'GET' :

		# J'initialise la valeur de l'argument "k_org_moa".
		try :
			v_org_moa = TMoa.objects.get(
				pk = TUtilisateur.objects.get(pk = request.user.pk).id_org,
				peu_doss = True,
				en_act_doss = True
			).pk
		except :
			v_org_moa = None

		# J'instancie un objet "formulaire".
		f_ch_doss = ChoisirDossier(prefix = 'ChoisirDossier', k_org_moa = v_org_moa)

		# Je prépare le tableau des dossiers.
		if v_org_moa :
			qs_doss = obt_doss_regr(v_org_moa)
		else :
			qs_doss = TDossier.objects.all()

		# 6880
		# qs_doss = TDossier.objects.custom_filter(
		# 	remove_completed = True, pk__in = qs_doss.values_list('pk', flat = True)
		# ) # Retrait des dossiers soldés
		# 6880: dans le cas ou on change d'avis...:
		# ces filtres sont a mis en regard des valeurs initiales de ChoisirDossier
		from django.db.models import Q
		selected = Q()
		selected.add(Q(id_av__int_av__icontains='Soldé'), Q.OR)
		selected.add(Q(id_av__int_av__icontains='Terminé'), Q.OR)
		selected.add(Q(id_av__int_av__icontains='Abandonné'), Q.OR)
		qs_doss = qs_doss.exclude(selected)



		t_doss = [{
			'pk' : d.pk,
			'num_doss' : d,
			'int_doss' : d.get_int_doss(),
			'n_org' : d.id_org_moa,
			'dt_delib_moa_doss' : dt_fr(d.dt_delib_moa_doss) or '-',
			'tx_engag_doss' : d.get_view_object().tx_engag_doss,
			'tx_real_doss' : d.get_view_object().tx_real_doss,
			'montant_dossier' : d.mont_doss,
		} for d in qs_doss]

		# J'affiche le template.
		output = render(
			request,
			'./gestion_dossiers/ch_doss.html',
			{ 'f_ch_doss' : init_f(f_ch_doss), 't_doss' : t_doss, 'title' : 'Choisir un dossier' }
		)

	else :
		if 'action' in request.GET :

			# Je stocke la valeur du paramètre "GET" dont la clé est "action".
			get_action = request.GET['action']

			# Je traite le cas où je dois alimenter les listes déroulantes des axes, des sous-axes, des actions et des
			# types de dossiers.
			if get_action == 'alimenter-listes' :

				# J'affiche ou je cache certaines listes déroulantes selon les données qu'elles contiennent.
				output = HttpResponse(json.dumps(alim_ld(request)), content_type = 'application/json')

			# Je traite le cas où je dois filtrer les dossiers selon certains critères.
			if get_action == 'filtrer-dossiers' :

				# Je prépare le tableau des dossiers filtrés.
				t_doss = [(
					d.num_doss,
					d.get_int_doss(),
					d.id_org_moa.n_org,
					dt_fr(d.dt_delib_moa_doss) or '-',
					str(d.get_view_object().tx_engag_doss),
					str(d.get_view_object().tx_real_doss),
					str(d.mont_doss),
					'<a href="{0}" class="consult-icon pull-right" title="Consulter le dossier"></a>'.format(
						reverse('cons_doss', args = [d.pk])
					)
				) for d in filtr_doss(request)]

				# J'envoie le tableau des dossiers filtrés.
				output = HttpResponse(
					json.dumps({ 'success' : { 'datatable' : t_doss}}), content_type = 'application/json'
				)

	return output

'''
Cette vue permet d'afficher la page de consultation d'un dossier ou de traiter une action.
request : Objet requête
_d : Identifiant d'un dossier
'''
@verif_acc
@csrf_exempt
def cons_doss(request, _d) :

	# Imports
	from app.forms.gestion_dossiers import AjouterPrestataire
	from app.forms.gestion_dossiers import ChoisirPrestation
	from app.forms.gestion_dossiers import GererArrete
	from app.forms.gestion_dossiers import GererDemandeVersement
	from app.forms.gestion_dossiers import GererDossier_Reglementations
	from app.forms.gestion_dossiers import GererFacture
	from app.forms.gestion_dossiers import GererFicheVie
	from app.forms.gestion_dossiers import GererFinancement
	from app.forms.gestion_dossiers import GererOrdreService
	from app.forms.gestion_dossiers import GererPrestation
	from app.forms.gestion_dossiers import GererPhoto
	from app.forms.gestion_dossiers import RedistribuerPrestation
	from app.forms.gestion_dossiers import PrintDoss
	from app.functions import dt_fr
	from app.functions import gen_f_ajout_aven
	from app.functions import ger_droits
	from app.functions import init_f
	from app.functions import init_fm
	from app.functions import init_pg_cons
	from app.functions import obt_mont
	from app.functions import obt_pourc
	from app.functions import rempl_fich_log
	from app.functions import suppr
	from app.models import TArretesDossier
	from app.models import TDossier
	from app.models import TDossierGeom
	from app.models import TFacture
	from app.models import TPhoto
	from app.models import TPrestation
	from app.models import TPrestationsDossier
	from app.models import TTypesGeomTypeDossier
	from app.models import VDemandeVersement
	from app.models import VFinancement
	from app.models import VPrestation
	from app.models import VSuiviDossier
	from app.models import VSuiviPrestationsDossier
	from django.contrib.gis import geos
	from django.core.urlresolvers import reverse
	from django.http import HttpResponse
	from django.shortcuts import get_object_or_404
	from django.shortcuts import render
	from django.template.context_processors import csrf
	from styx.settings import T_DONN_BDD_STR
	from styx.settings import MEDIA_URL
	import json

	output = HttpResponse()

	# Je vérifie l'existence d'un objet TDossier.
	o_doss = get_object_or_404(TDossier, pk = _d)

	if request.method == 'GET' :

		# Je vérifie le droit de lecture.
		ger_droits(request.user, o_doss)

		# Je désigne l'onglet actif du navigateur à onglets relatif à un dossier.
		if 'tab_doss' not in request.session :
			request.session['tab_doss'] = '#ong_doss'

		# Je définis si le montant du dossier est en HT ou en TTC.
		ht_ou_ttc = 'HT'
		if o_doss.est_ttc_doss == True :
			ht_ou_ttc = 'TTC'

		# Je pointe vers l'objet VSuiviDossier.
		o_suivi_doss = VSuiviDossier.objects.get(id_doss = o_doss.pk)

		# Obtention des différents tableaux déployés dans chaque onglet
		doss_fam = o_doss.get_doss_fam()
		fdvs = o_doss.get_recap_fdvs()
		fins = o_doss.get_recap_fins()
		prss = o_doss.get_recap_prss()
		facs = o_doss.get_recap_facs()
		ddvs = o_doss.get_recap_ddvs()
		arrs = o_doss.get_recap_arrs()

		# Je stocke le jeu de données des prestations du maître d'ouvrage du dossier.
		qs_prest_moa_doss = TPrestation.objects.filter(
			tprestationsdossier__id_doss__id_org_moa = o_doss.id_org_moa
		).exclude(tprestationsdossier__id_doss = o_doss).distinct()

		# J'initialise le tableau des prestations du maître d'ouvrage du dossier.
		t_prest_moa_doss = []
		for pmd in qs_prest_moa_doss :

			# Je vérifie si je peux choisir la prestation courante pour une redistribution.
			peut_ch = True
			for dp in TPrestationsDossier.objects.filter(id_prest = pmd.pk) :
				ht_ou_ttc_dp = 'HT'
				if dp.id_doss.est_ttc_doss == True :
					ht_ou_ttc_dp = 'TTC'
				if ht_ou_ttc != ht_ou_ttc_dp :
					peut_ch = False

			# J'initialise le contenu de la dernière balise <td/>.
			dern_td = ''
			if peut_ch == True :
				dern_td = '''
				<span action="?action=afficher-form-redistribution-prestation&prestation={0}"
				class="choose-icon pointer pull-right" title="Choisir la prestation"></span>
				'''.format(pmd.pk)

			lg = '''
			<tr>
				<td>{0}</td>
				<td>{1}</td>
				<td>{2}</td>
				<td>{3}</td>
				<td>{4}</td>
				<td>{5}</td>
			</tr>
			'''.format(
				pmd.id_org_prest.n_org,
				pmd.int_prest,
				dt_fr(pmd.dt_notif_prest) or '-',
				obt_mont(VPrestation.objects.get(pk = pmd.pk).mont_prest),
				', '.join([(dp.id_doss.num_doss) for dp in TPrestationsDossier.objects.filter(id_prest = pmd).order_by(
					'id_doss'
				)]),
				dern_td
			)
			t_prest_moa_doss.append(lg)

		# Je stocke le jeu de données des photos du dossier.
		qs_ph = TPhoto.objects.filter(id_doss = o_doss)

		# J'initialise le tableau des photos du dossier.
		t_ph = [{
			'id_doss' : p.id_doss.pk,
			'chem_ph' : p.chem_ph,
			'int_ph' : p.int_ph,
			'int_ppv_ph' : p.id_ppv_ph,
			'dt_pv_ph' : dt_fr(p.dt_pv_ph) or '-',
			'pk' : p.pk
		} for p in qs_ph]

		# Je prépare la fenêtre modale relative au diaporama des photos du dossier.
		t_ph_diap = { 'li' : [], 'div' : [] }
		for index, p in enumerate(qs_ph) :

			t_attrs_li = ['data-target="#my-carousel"', 'data-slide-to="{0}"'.format(index)]
			if index == 0 :
				t_attrs_li.append('class="active"')
			li = '<li {0}></li>'.format(' '.join(t_attrs_li))

			class_div = 'item'
			if index == 0 :
				class_div += ' active'
			div = '''
			<div class="{0}">
				<img src="{1}">
			</div>
			'''.format(class_div, MEDIA_URL + str(p.chem_ph))

			t_ph_diap['li'].append(li)
			t_ph_diap['div'].append(div)

		# J'initialise le tableau des géométries du dossier.
		qs_geom_doss = TDossierGeom.objects.filter(id_doss = o_doss)
		t_geom_doss = []
		for g in qs_geom_doss :
			if g.geom_pol :
				la_geom = geos.GEOSGeometry(g.geom_pol)
			if g.geom_lin :
				la_geom = geos.GEOSGeometry(g.geom_lin)
			if g.geom_pct :
				la_geom = geos.GEOSGeometry(g.geom_pct)
			t_geom_doss.append(la_geom.geojson)

		# J'initialise les types de géométries autorisés.
		qs_type_geom_doss = TTypesGeomTypeDossier.objects.filter(id_type_doss = o_doss.id_type_doss)
		t_types_geom_doss = [tg.id_type_geom.int_type_geom for tg in qs_type_geom_doss]

		# J'instancie des objets "formulaire".
		f_ajout_fin = GererFinancement(prefix = 'GererFinancement', k_doss = o_doss)
		f_ajout_prest = GererPrestation(prefix = 'GererPrestation', k_doss = o_doss)
		f_ch_prest = ChoisirPrestation(prefix = 'ChoisirPrestation', k_org_moa = o_doss.id_org_moa)
		f_ajout_fact = GererFacture(prefix = 'GererFacture', k_doss = o_doss)
		f_ajout_ddv = GererDemandeVersement(prefix = 'GererDemandeVersement', k_doss = o_doss)
		f_ajout_arr = GererArrete(prefix = 'GererArrete', k_doss = o_doss)
		f_modif_doss_regl = GererDossier_Reglementations(prefix = 'GererDossier', instance = o_doss)
		f_ajout_ph = GererPhoto(prefix = 'GererPhoto', k_doss = o_doss)
		f_ajout_org_prest = AjouterPrestataire(prefix = 'AjouterPrestataire')
		f_ajout_fdv = GererFicheVie(k_doss = o_doss)

		# J'initialise le gabarit de chaque champ de chaque formulaire.
		t_ajout_fin = init_f(f_ajout_fin)
		t_ajout_prest = init_f(f_ajout_prest)
		t_ch_prest = init_f(f_ch_prest)
		t_ajout_fact = init_f(f_ajout_fact)
		t_ajout_ddv = init_f(f_ajout_ddv)
		t_ajout_arr = init_f(f_ajout_arr)
		t_ajout_ph = init_f(f_ajout_ph)
		t_ajout_org_prest = init_f(f_ajout_org_prest)

		# Je déclare un tableau qui stocke le contenu de certaines fenêtres modales.
		t_cont_fm = {
			'ajout_arr' : '''
			<form action="{0}" enctype="multipart/form-data" name="f_ajout_arr" method="post" onsubmit="soum_f(event)">
				<input name="csrfmiddlewaretoken" type="hidden" value="{1}">
				<div class="row">
					<div class="col-sm-6">{2}</div>
					<div class="col-sm-6">{3}</div>
				</div>
				{4}
				{5}
				{6}
				{7}
				{8}
				{9}
				<button class="center-block green-btn my-btn" type="submit">Valider</button>
				<div class="form-remark">
					<ul class="my-list-style">
						<li>Le numéro de l'arrêté, la date de signature de l'arrêté, la date limite d'enclenchement des
						travaux et le fichier scanné de l'arrêté sont obligatoires si et seulement si l'avancement de
						l'arrêté est « Validé ».</li>
					</ul>
				</div>
			</form>
			'''.format(
				reverse('ajout_arr'),
				csrf(request)['csrf_token'],
				t_ajout_arr['za_num_doss'],
				t_ajout_arr['zl_type_decl'],
				t_ajout_arr['id_type_av_arr'],
				t_ajout_arr['num_arr'],
				t_ajout_arr['dt_sign_arr'],
				t_ajout_arr['dt_lim_encl_trav_arr'],
				t_ajout_arr['chem_pj_arr'],
				t_ajout_arr['comm_arr']
			),
			'ajout_ddv' : '''
			<form action="{0}" enctype="multipart/form-data" name="f_ger_ddv" method="post" onsubmit="soum_f(event)">
				<input name="csrfmiddlewaretoken" type="hidden" value="{1}">
				{2}
				{3}
				{4}
				{5}
				{6}
				<div class="row">
					<div class="col-md-6">{7}</div>
					<div class="col-md-6">{8}</div>
				</div>
				<div class="row">
					<div class="col-sm-6">{9}</div>
					<div class="col-sm-6">{10}</div>
				</div>
				<div class="row">
					<div class="col-sm-6">{11}</div>
					<div class="col-sm-6">{12}</div>
				</div>
				<div class="row">
					<div class="col-sm-6">{13}</div>
					<div class="col-sm-6">{14}</div>
				</div>
				{15}
				{16}
				<button class="center-block green-btn my-btn" type="submit">Valider</button>
				<div class="form-remark">
					<ul class="my-list-style">
						<li>Il est impossible de renseigner un numéro de bordereau et un numéro de titre de recette
						tant qu'une date de versement n'a pas été renseignée.</li>
					</ul>
				</div>
			</form>
			'''.format(
				reverse('ajout_ddv'),
				csrf(request)['csrf_token'],
				t_ajout_ddv['za_num_doss'],
				t_ajout_ddv['zl_fin'],
				t_ajout_ddv['id_type_vers'],
				t_ajout_ddv['int_ddv'],
				t_ajout_ddv['cbsm_fact'],
				t_ajout_ddv['mont_ht_ddv'],
				t_ajout_ddv['mont_ttc_ddv'],
				t_ajout_ddv['dt_ddv'],
				t_ajout_ddv['dt_vers_ddv'],
				t_ajout_ddv['num_bord_ddv'],
				t_ajout_ddv['num_titre_rec_ddv'],
				t_ajout_ddv['mont_ht_verse_ddv'],
				t_ajout_ddv['mont_ttc_verse_ddv'],
				t_ajout_ddv['chem_pj_ddv'],
				t_ajout_ddv['comm_ddv'],
			),
			'ajout_fact' : '''
			<form action="{}" enctype="multipart/form-data" name="f_ajout_fact" method="post" onsubmit="soum_f(event)">
				<input name="csrfmiddlewaretoken" type="hidden" value="{}">
				{}
				{}
				{}
				{}
				<div class="row">
					<div class="col-sm-6">{}</div>
					<div class="col-sm-6">{}</div>
				</div>
				{}
				<div class="row">
					<div class="col-sm-6">{}</div>
					<div class="col-sm-6">{}</div>
				</div>
				{}
				{}
				{}
				<button class="center-block green-btn my-btn" type="submit">Valider</button>
			</form>
			'''.format(
				reverse('ajout_fact'),
				csrf(request)['csrf_token'],
				t_ajout_fact['za_num_doss'],
				t_ajout_fact['zl_prest'],
				t_ajout_fact['num_fact'],
				t_ajout_fact['dt_rec_fact'],
				t_ajout_fact['mont_ht_fact'],
				t_ajout_fact['mont_ttc_fact'],
				t_ajout_fact['dt_mand_moa_fact'],
				t_ajout_fact['num_bord_fact'],
				t_ajout_fact['num_mandat_fact'],
				t_ajout_fact['zl_suivi_fact'],
				t_ajout_fact['chem_pj_fact'],
				t_ajout_fact['comm_fact']
			),
			'ajout_fin' : '''
			<form action="{0}" enctype="multipart/form-data" name="f_ajout_fin" method="post" onsubmit="soum_f(event)">
				<input name="csrfmiddlewaretoken" type="hidden" value="{1}">
				{2}
				{3}
				{4}
				<div class="row">
					<div class="col-md-6">{5}</div>
					<div class="col-md-6">{6}</div>
				</div>
				{7}
				{8}
				<div class="row">
					<div class="col-sm-6">{9}</div>
					<div class="col-sm-6">{10}</div>
				</div>
				{11}
				{12}
				<div class="row">
					<div class="col-sm-6">{13}</div>
					<div class="col-sm-6">{14}</div>
				</div>
				{15}
				{16}
				{17}
				<button class="center-block green-btn my-btn" type="submit">Valider</button>
				<div class="form-remark">
					<ul class="my-list-style">
						<li>Le montant de l'assiette éligible de la subvention et le pourcentage de l'assiette éligible
						peuvent ne pas être renseignés, mais si l'un d'entre eux est renseigné, alors l'autre doit
						l'être également.</li>
						<li>La date de début d'éligibilité doit être renseignée si une durée de validité de l'aide ou
						une durée de prorogation est renseignée.</li>
						<li>Si la date limite du début de l'opération est renseignée, alors seules les options "Oui" et
						"Non" de la liste déroulante suivant le champ en question seront valides. Dans le cas
						contraire, seule l'option "Sans objet" sera valide.</li>
						<li>Il est impossible de renseigner un pourcentage de réalisation des travaux tant que le
						premier acompte n'est pas payé en fonction de celui-ci.</li>
					</ul>
				</div>
			</form>
			'''.format(
				reverse('ajout_fin'),
				csrf(request)['csrf_token'],
				t_ajout_fin['za_num_doss'],
				t_ajout_fin['id_org_fin'],
				t_ajout_fin['num_arr_fin'],
				t_ajout_fin['mont_elig_fin'],
				t_ajout_fin['pourc_elig_fin'],
				t_ajout_fin['mont_part_fin'],
				t_ajout_fin['dt_deb_elig_fin'],
				t_ajout_fin['duree_valid_fin'],
				t_ajout_fin['duree_pror_fin'],
				t_ajout_fin['dt_lim_deb_oper_fin'],
				t_ajout_fin['a_inf_fin'],
				t_ajout_fin['dt_lim_prem_ac_fin'],
				t_ajout_fin['id_paiem_prem_ac'],
				t_ajout_fin['pourc_real_fin'],
				t_ajout_fin['chem_pj_fin'],
				t_ajout_fin['comm_fin']
			),
			'ajout_org_prest' : '''
			<form action="{0}" name="f_ajout_org_prest" method="post" onsubmit="soum_f(event)">
				<input name="csrfmiddlewaretoken" type="hidden" value="{1}">
				<div class="row">
					<div class="col-sm-6">{2}</div>
					<div class="col-sm-6">{3}</div>
				</div>
				{4}
				{5}
				<button class="center-block green-btn my-btn" type="submit">Valider</button>
			</form>
			'''.format(
				reverse('ajout_org_prest'),
				csrf(request)['csrf_token'],
				t_ajout_org_prest['n_org'],
				t_ajout_org_prest['siret_org_prest'],
				t_ajout_org_prest['num_dep'],
				t_ajout_org_prest['comm_org']
			),
			'ajout_prest' : '''
			<form>{}</form>
			<div id="za_nvelle_prest">
				<form action="{}" enctype="multipart/form-data" name="f_ajout_prest" method="post" onsubmit="soum_f(event)">
					<input name="csrfmiddlewaretoken" type="hidden" value="{}">
					{}
					<div class="row">
						<div class="col-sm-6">{}</div>
						<div class="col-sm-6" style="margin: 20px 0;">
							<span class="add-company-icon icon-link" id="bt_ajout_org_prest">Ajouter un prestataire</span>
						</div>
					</div>
					{}
					{}
					{}
					{}
					<div class="row">
						<div class="col-sm-6">{}</div>
						<div class="col-sm-6">{}</div>
					</div>
					{}
					{}
					{}
					<button class="center-block green-btn my-btn" type="submit">Valider</button>
				</form>
			</div>
			<div id="za_red_prest_etape_1" style="display: none;">
				<form action="?action=filtrer-prestations" name="f_ch_prest" method="post">
					<input name="csrfmiddlewaretoken" type="hidden" value="{}">
					<fieldset class="my-fieldset" style="padding-bottom: 0;">
						<legend>Rechercher par</legend>
						<div>
							{}
							{}
							{}
						</div>
					</fieldset>
				</form>
				<div class="br"></div>
				<span class="my-label">
					<span class="u">Étape 1 :</span> Choisir une prestation dont le montant total est en {}
				</span>
				<div class="my-table" id="t_ch_prest">
					<table>
						<thead>
							<tr>
								<th>Prestataire</th>
								<th>Intitulé de la prestation</th>
								<th>Date de notification</th>
								<th>Montant (en €)</th>
								<th>Dossier(s)</th>
								<th></th>
							</tr>
						</thead>
						<tbody>{}</tbody>
					</table>
				</div>
			</div>
			'''.format(
				t_ajout_prest['rb_prest_exist'],
				reverse('ajout_prest'),
				csrf(request)['csrf_token'],
				t_ajout_prest['za_num_doss'],
				t_ajout_prest['zs_siret_org_prest'],
				t_ajout_prest['int_prest'],
				t_ajout_prest['ref_prest'],
				t_ajout_prest['zs_mont_prest'],
				t_ajout_prest['zs_duree_prest'],
				t_ajout_prest['dt_notif_prest'],
				t_ajout_prest['dt_fin_prest'],
				t_ajout_prest['id_nat_prest'],
				t_ajout_prest['chem_pj_prest'],
				t_ajout_prest['comm_prest'],
				csrf(request)['csrf_token'],
				t_ch_prest['zl_progr'],
				t_ch_prest['za_org_moa'],
				t_ch_prest['zl_org_prest'],
				ht_ou_ttc,
				'\n'.join(t_prest_moa_doss)
			),
			'ajout_ph' : '''
			<form action="{0}" enctype="multipart/form-data" name="f_ajout_ph" method="post" onsubmit="soum_f(event)">
				<input name="csrfmiddlewaretoken" type="hidden" value="{1}">
				{2}
				<div class="row">
					<div class="col-sm-6">{3}</div>
					<div class="col-sm-6">{4}</div>
				</div>
				{5}
				{6}
				{7}
				<button class="center-block green-btn my-btn" type="submit">Valider</button>
			</form>
			'''.format(
				reverse('ajout_ph'),
				csrf(request)['csrf_token'],
				t_ajout_ph['za_num_doss'],
				t_ajout_ph['int_ph'],
				t_ajout_ph['descr_ph'],
				t_ajout_ph['id_ppv_ph'],
				t_ajout_ph['dt_pv_ph'],
				t_ajout_ph['chem_ph']
			),
			'lanc_diap' : '''
			<div class="carousel my-carousel slide" data-interval="false" data-ride="carousel" id="my-carousel">
				<ol class="carousel-indicators">{0}</ol>
				<div class="carousel-inner" role="listbox">{1}</div>
				<a href="#my-carousel" class="left carousel-control" data-slide="prev" role="button">
					<span class="glyphicon glyphicon-chevron-left" aria-hidden="true"></span>
				</a>
				<a href="#my-carousel" class="right carousel-control" data-slide="next" role="button">
					<span class="glyphicon glyphicon-chevron-right" aria-hidden="true"></span>
				</a>
			</div>
			'''.format('\n'.join(t_ph_diap['li']), '\n'.join(t_ph_diap['div'])),
			'impr_doss' : '''
			<form action="{}" method="post" name="print_doss" target="_blank">
				<input name="csrfmiddlewaretoken" type="hidden" value="{}">
				{}
				<div class="br"></div>
				<input type="hidden" name="action" value="print_doss"/>
				<input class="center-block green-btn my-btn" type="submit" name="confirm" value="Imprimer"/>
			</form>
			'''.format(
				reverse('impr_doss', args = [o_doss.pk]),
				csrf(request)['csrf_token'],
				'<br>'.join(['{} {}'.format(f, f.label) for f in PrintDoss()])
			)
		}

		# Je déclare un tableau de fenêtres modales.
		t_fm = [
			init_fm('ajout_arr', 'Ajouter un arrêté', t_cont_fm['ajout_arr']),
			init_fm('ajout_fin', 'Ajouter un organisme financier', t_cont_fm['ajout_fin']),
			init_fm('ajout_ph', 'Ajouter une photo', t_cont_fm['ajout_ph']),
			init_fm('gerfdv', 'Ajouter un événement', f_ajout_fdv.get_form(rq = request)),
			init_fm('geros', 'Ajouter un ordre de service'),
			init_fm('impr_doss', 'Imprimer un dossier - Sélection des onglets', t_cont_fm['impr_doss']),
			init_fm('modif_carto', 'Modifier un dossier'),
			init_fm('modif_doss_regl', 'Modifier un dossier'),
			init_fm('suppr_doss', 'Supprimer un dossier'),
			init_fm('suppr_fdv', 'Supprimer un événement')
		]

		# Je complète le tableau de fenêtres modales dans le cas où le dossier n'est pas en projet.
		if o_doss.id_av.int_av != T_DONN_BDD_STR['AV_EP'] :
			t_fm += [
				init_fm('ajout_aven', 'Ajouter un avenant'),
				init_fm('ajout_fact', 'Ajouter une facture', t_cont_fm['ajout_fact']),
				init_fm('ajout_org_prest', 'Ajouter un prestataire', t_cont_fm['ajout_org_prest']),
				init_fm('ajout_prest', 'Ajouter/relier une prestation', t_cont_fm['ajout_prest']),
				init_fm('ger_ddv', 'Ajouter une demande de versement', t_cont_fm['ajout_ddv']),
			]

		# Je complète le tableau de fenêtres modales dans le cas où le dossier comporte des photos.
		if len(t_ph) > 0 :
			t_fm += [
				init_fm('cons_ph', 'Consulter une photo'),
				init_fm('lanc_diap', 'Lancer le diaporama', t_cont_fm['lanc_diap']),
				init_fm('suppr_ph', 'Supprimer une photo')
			]

		# J'affiche le template.
		output = render(
			request,
			'./gestion_dossiers/cons_doss.html',
			{
				'AV_EP' : T_DONN_BDD_STR['AV_EP'],
				'd' : o_doss,
				'f_modif_doss_regl' : init_f(f_modif_doss_regl),
				'forbidden' : ger_droits(request.user, o_doss, False, False),
				'ht_ou_ttc' : ht_ou_ttc,
				'mont_ddv_sum' : ddvs['mnt'],
				'mont_ddv_sum_str' : obt_mont(ddvs['mnt']),
				'mont_doss' : obt_mont(o_doss.mont_doss),
				'mont_fact_sum' : facs['mnt'],
				'mont_fact_sum_str' : obt_mont(facs['mnt']),
				'mont_rae' : obt_mont(o_suivi_doss.mont_rae),
				'mont_suppl_doss' : obt_mont(o_doss.mont_suppl_doss),
				't_arr' : arrs,
				't_attrs_doss' : init_pg_cons(o_doss.get_attrs()),
				't_ddv' : ddvs['tbl'],
				't_doss_fam' : doss_fam,
				't_fact' : facs['tbl'],
				't_fdvs' : fdvs,
				't_fin' : fins,
				't_fm' : t_fm,
				't_geom_doss' : t_geom_doss,
				't_ph' : t_ph,
				't_ph_length' : len(t_ph),
				't_prest' : prss['tbl'],
				't_prest_length' : len(prss['tbl']),
				't_prest_sum' : prss['mnts'],
				't_types_geom_doss' : t_types_geom_doss,
				'title' : 'Consulter un dossier'
			}
		)

		# Je supprime la variable de session liée au navigateur à onglets relatif à un dossier.
		del request.session['tab_doss']

	else :
		if 'action' in request.GET :

			# Je stocke la valeur du paramètre "GET" dont la clé est "action".
			get_action = request.GET['action']

			# Je traite le cas où je dois supprimer une photo.
			if get_action == 'supprimer-photo' :
				if request.GET['photo'] :
					output = HttpResponse(suppr(reverse('suppr_ph', args = [request.GET['photo']])))

			# Je traite le cas où je dois consulter une photo.
			if get_action == 'consulter-photo' :
				if request.GET['photo'] :

					# Je vérifie l'existence d'un objet TPhoto.
					o_ph = get_object_or_404(TPhoto, pk = request.GET['photo'])

					# Je vérifie le droit de lecture.
					ger_droits(request.user, o_ph.id_doss)

					# Je prépare le volet de consultation de la photo.
					t_attrs_ph = {
						'id_doss' : { 'label' : 'Numéro du dossier', 'value' : o_ph.id_doss },
						'int_ph' : { 'label' : 'Intitulé de la photo', 'value' : o_ph.int_ph },
						'descr_ph' : { 'label' : 'Description', 'value' : o_ph.descr_ph or '' },
						'id_ppv_ph' : { 'label' : 'Période de prise de vue', 'value' : o_ph.id_ppv_ph },
						'dt_pv_ph' : { 'label' : 'Date de prise de vue', 'value' : dt_fr(o_ph.dt_pv_ph) or '' }
					}

					t_attrs_ph = init_pg_cons(t_attrs_ph)

					output = HttpResponse(
					'''
					<div style="margin-bottom: -20px;">
						<div class="attribute-wrapper">
							<span class="attribute-label">Visualisation</span>
							<img src="{0}{1}" style="display: block; margin: 0 auto; max-width: 100%;">
						</div>
						{2}
						<div class="row">
							<div class="col-sm-6">{3}</div>
							<div class="col-sm-6">{4}</div>
						</div>
						<div class="row">
							<div class="col-sm-6">{5}</div>
							<div class="col-sm-6">{6}</div>
						</div>
					</div>
					'''.format(
						MEDIA_URL,
						o_ph.chem_ph,
						t_attrs_ph['id_doss'],
						t_attrs_ph['int_ph'],
						t_attrs_ph['descr_ph'],
						t_attrs_ph['id_ppv_ph'],
						t_attrs_ph['dt_pv_ph']
					))

			# Je traite le cas où je dois modifier le commentaire de l'onglet "Réglementations".
			if get_action == 'modifier-reglementation' :

				# Je vérifie le droit d'écriture.
				ger_droits(request.user, o_doss, False)

				# Je soumets le formulaire.
				f_modif_doss_regl = GererDossier_Reglementations(
					request.POST, instance = o_doss, prefix = 'GererDossier'
				)

				if f_modif_doss_regl.is_valid() :

					# Je modifie l'instance TDossier.
					o_doss_regl_modif = f_modif_doss_regl.save(commit = False)
					o_doss_regl_modif.save()

					# J'affiche le message de succès.
					output = HttpResponse(
						json.dumps({ 'success' : {
							'message' : 'Le dossier {0} a été mis à jour avec succès.'.format(o_doss_regl_modif),
							'redirect' : reverse('cons_doss', args = [o_doss_regl_modif.pk])
						}}),
						content_type = 'application/json'
					)

					# Je renseigne l'onglet actif après rechargement de la page.
					request.session['tab_doss'] = '#ong_arr'

					# Je complète le fichier log.
					rempl_fich_log([
						request.user.pk,
						request.user,
						o_doss_regl_modif.pk,
						o_doss_regl_modif,
						'U',
						'Réglementation d\'un dossier',
						o_doss_regl_modif.pk
					])

				else :

					# J'affiche les erreurs.
					t_err = {}
					for k, v in f_modif_doss_regl.errors.items() :
						t_err['GererDossier-{0}'.format(k)] = v
					output = HttpResponse(json.dumps(t_err), content_type = 'application/json')

			# Je traite le cas où je dois afficher le formulaire d'ajout d'un avenant.
			if get_action == 'afficher-form-avenant' :
				if request.GET['prestation'] :

					# Je vérifie l'existence d'un objet TPrestationsDossier.
					o_prest_doss = get_object_or_404(TPrestationsDossier, pk = request.GET['prestation'])

					output = HttpResponse(gen_f_ajout_aven(request, o_prest_doss, reverse('ajout_aven_racc')))

			# Je traite le cas où je dois afficher le formulaire d'ajout d'un ordre de service.
			if get_action == 'afficher-form-os' :
				if request.GET['prestation'] :
					o_pd = get_object_or_404(TPrestationsDossier, pk = request.GET['prestation'])
					output = HttpResponse(GererOrdreService(k_pd = o_pd, prefix = 'GererOrdreService').get_form(
						racc = True, rq = request
					))

			# Je traite le cas où je dois filtrer les prestations.
			if get_action == 'filtrer-prestations' :

				# Je soumets le formulaire.
				f_ch_prest = ChoisirPrestation(request.POST, prefix = 'ChoisirPrestation')

				if f_ch_prest.is_valid() :

					# Je récupère les données du formulaire valide.
					cleaned_data = f_ch_prest.cleaned_data
					v_progr = cleaned_data.get('zl_progr')
					v_org_prest = cleaned_data.get('zl_org_prest')

					# J'initialise les conditions de la requête.
					t_sql = { 'and' : {}, 'or' : [] }
					if v_progr :
						t_sql['and']['tprestationsdossier__id_doss__id_progr'] = v_progr
					t_sql['and']['tprestationsdossier__id_doss__id_org_moa__id_org_moa'] = o_doss.id_org_moa.pk
					if v_org_prest :
						t_sql['and']['id_org_prest'] = v_org_prest

					# Je stocke le jeu de données des prestations filtrées du maître d'ouvrage du dossier.
					qs_prest_moa_doss_filtr = TPrestation.objects.filter(**t_sql['and']).exclude(
						tprestationsdossier__id_doss = o_doss
					).distinct().order_by('id_org_prest', 'int_prest')

					# Je définis si le montant du dossier est en HT ou en TTC.
					ht_ou_ttc = 'HT'
					if o_doss.est_ttc_doss == True :
						ht_ou_ttc = 'TTC'

					# J'initialise le tableau des prestations filtrées du maître d'ouvrage du dossier.
					t_prest_moa_doss_filtr = []
					for pmd in qs_prest_moa_doss_filtr :

						# Je vérifie si je peux choisir la prestation courante pour une redistribution.
						peut_ch = True
						for dp in TPrestationsDossier.objects.filter(id_prest = pmd.pk) :
							ht_ou_ttc_dp = 'HT'
							if dp.id_doss.est_ttc_doss == True :
								ht_ou_ttc_dp = 'TTC'
							if ht_ou_ttc != ht_ou_ttc_dp :
								peut_ch = False

						# J'initialise le contenu de la dernière balise <td/>.
						dern_td = ''
						if peut_ch == True :
							dern_td = '''
							<span action="?action=afficher-form-redistribution-prestation&prestation={0}"
							class="choose-icon pointer pull-right" title="Choisir la prestation"></span>
							'''.format(pmd.pk)

						t_prest_moa_doss_filtr.append([
							pmd.id_org_prest.n_org,
							pmd.int_prest,
							dt_fr(pmd.dt_notif_prest),
							obt_mont(VPrestation.objects.get(pk = pmd.pk).mont_prest),
							', '.join([
								(dp.id_doss.num_doss) for dp in TPrestationsDossier.objects.filter(
									id_prest = pmd
								).order_by('id_doss')
							]),
							dern_td
						])

					# J'envoie le tableau des prestations filtrées.
					output = HttpResponse(
						json.dumps({
							'success' : { 'datatable' : t_prest_moa_doss_filtr }
						}), content_type = 'application/json'
					)

				else :

					# J'affiche les erreurs.
					t_err = {}
					for k, v in f_ch_prest.errors.items() :
						t_err['ChoisirPrestation-{0}'.format(k)] = v
					output = HttpResponse(json.dumps(t_err), content_type = 'application/json')

			# Je traite le cas où je dois afficher le formulaire de redistribution des montants d'une prestation.
			if get_action == 'afficher-form-redistribution-prestation' :

				# J'initialise l'identifiant de la prestation (-1 étant un identifiant inexistant).
				v_prest = -1
				if request.GET['prestation'] :
					v_prest = request.GET['prestation']

				# Je vérifie l'existence d'un objet TPrestation.
				o_prest = get_object_or_404(TPrestation, pk = v_prest)

				t_lg = []
				for index, dp in enumerate(TPrestationsDossier.objects.filter(id_prest = o_prest)) :

					# Je pointe vers l'objet VSuiviPrestationsDossier.
					o_suivi_prest_doss = VSuiviPrestationsDossier.objects.get(pk = dp.pk)

					# Je pointe vers l'objet VSuiviDossier.
					o_suivi_doss = VSuiviDossier.objects.get(pk = dp.id_doss.pk)

					# J'instancie un objet "formulaire".
					f_red_prest = RedistribuerPrestation(
						instance = dp, prefix = 'RedistribuerPrestation{0}'.format(index)
					)

					# J'initialise le gabarit de chaque champ du formulaire.
					t_red_prest = init_f(f_red_prest)

					# J'initialise la somme des factures émises selon le mode de taxe du dossier.
					mont_fact_sum = o_suivi_prest_doss.mont_ht_fact_sum
					if o_doss.est_ttc_doss == True :
						mont_fact_sum = o_suivi_prest_doss.mont_ttc_fact_sum

					lg = '''
					<tr>
						<td class="b">{0}</td>
						<td>{1}</td>
						<td>{2}</td>
						<td>{3}</td>
						<td>{4}</td>
						<td>{5}</td>
					</tr>
					'''.format(
						dp.id_doss.num_doss,
						t_red_prest['mont_prest_doss'],
						obt_mont(o_suivi_prest_doss.mont_aven_sum),
						obt_mont(mont_fact_sum),
						obt_mont(o_suivi_doss.mont_rae),
						t_red_prest['duree_prest_doss']
					)

					# J'empile le tableau des lignes du tableau HTML des dossiers déjà reliés à la prestation que l'on
					# veut redistribuer.
					t_lg.append(lg)

				# Je pointe vers l'objet VSuiviDossier.
				o_suivi_doss = VSuiviDossier.objects.get(pk = o_doss.pk)

				# J'instancie un objet "formulaire".
				f_red_prest = RedistribuerPrestation(prefix = 'RedistribuerPrestation', k_doss = o_doss)

				# J'initialise le gabarit de chaque champ du formulaire.
				t_red_prest = init_f(f_red_prest)

				lg = '''
				<tr style="background-color: #F8B862;">
					<td class="b">{0}</td>
					<td>{1}</td>
					<td>0</td>
					<td>0</td>
					<td>{2}</td>
					<td>{3}</td>
				</tr>
				'''.format(
					o_doss.num_doss,
					t_red_prest['mont_prest_doss'],
					obt_mont(o_suivi_doss.mont_rae),
					t_red_prest['duree_prest_doss']
				)

				# J'empile le tableau des lignes du tableau HTML de redistribution d'une prestation en ajoutant la
				# ligne relative au dossier que l'on souhaite relier à la prestation choisie.
				t_lg.append(lg)

				# Je définis le mode de taxe du dossier.
				ht_ou_ttc = 'HT'
				if o_suivi_doss.est_ttc_doss == True :
					ht_ou_ttc = 'TTC'

				# J'envoie le formulaire de redistribution d'une prestation.
				output = HttpResponse(
					'''
					<div id="za_red_prest_etape_2" style="margin-top: 20px;">
						<span class="my-label">
							<span class="u">Étape 2 :</span> Redistribuer les montants dédiés de cette prestation
						</span>
						<form action="?action=relier-prestation" name="f_red_prest" method="post" onsubmit="soum_f(event)">
							<input name="csrfmiddlewaretoken" type="hidden" value="{0}">
							<div style="margin-bottom: 20px;">
								<div class="my-table" id="t_red_prest">
									<table>
										<thead>
											<tr>
												<th>N° du dossier</th>
												<th>Montant {1} de la prestation</th>
												<th>Somme {2} des avenants (en €)</th>
												<th>Somme {3} des factures émises (en €)</th>
												<th>Reste à engager {4} pour le dossier (en €)</th>
												<th>Durée de la prestation (en nombre de jours ouvrés)</th>
											</tr>
										</thead>
										<tbody>{5}</tbody>
									</table>
								</div>
							</div>
							<button class="center-block green-btn my-btn" type="submit">Valider</button>
						</form>
					</div>
					'''.format(
						csrf(request)['csrf_token'],
						ht_ou_ttc,
						ht_ou_ttc,
						ht_ou_ttc,
						ht_ou_ttc,
						'\n'.join(t_lg))
				)

				# Je stocke en mémoire l'identifiant de la prestation valide.
				request.session['relier-prestation'] = o_prest.pk

			# Je traite le cas où je dois relier une prestation à un autre dossier.
			if get_action == 'relier-prestation' :

				# J'initialise l'identifiant de la prestation (-1 étant un identifiant inexistant).
				v_prest = -1
				if 'relier-prestation' in request.session :
					v_prest = request.session['relier-prestation']

				# Je vérifie l'existence d'un objet TPrestation.
				o_prest = get_object_or_404(TPrestation, pk = v_prest)

				# Je vérifie le droit d'écriture.
				ger_droits(request.user, o_doss, False)

				# J'initialise le tableau des erreurs tous formulaires confondus.
				t_err = {}

				# J'initialise le tableau des instances TPrestationsDossier à créer ou à modifier.
				t_inst_prest_doss = []

				for index, dp in enumerate(TPrestationsDossier.objects.filter(id_prest = o_prest)) :

					# Je soumets le formulaire.
					f_red_prest = RedistribuerPrestation(
						request.POST, instance = dp, prefix = 'RedistribuerPrestation{0}'.format(index)
					)

					if f_red_prest.is_valid() :

						# J'empile le tableau des instances TPrestationsDossier.
						t_inst_prest_doss.append(f_red_prest.save(commit = False))

					else :

						# J'empile le tableau des erreurs.
						for k, v in f_red_prest.errors.items() :
							t_err['RedistribuerPrestation{0}-{1}'.format(index, k)] = v

				# Je soumets le formulaire.
				f_red_prest = RedistribuerPrestation(
					request.POST, prefix = 'RedistribuerPrestation', k_doss = o_doss, k_prest = o_prest
				)

				if f_red_prest.is_valid() :

					# J'empile le tableau des instances TPrestationsDossier.
					t_inst_prest_doss.append(f_red_prest.save(commit = False))

				else :

					# J'empile le tableau des erreurs.
					for k, v in f_red_prest.errors.items() :
						t_err['RedistribuerPrestation-{0}'.format(k)] = v

				if len(t_err) > 0 :

					# J'affiche les erreurs.
					output = HttpResponse(json.dumps(t_err), content_type = 'application/json')

				else :

					# Je créé ou modifie chaque instance TPrestationsDossier.
					for i in range(0, len(t_inst_prest_doss)) :
						t_inst_prest_doss[i].save()

					# J'affiche le message de succès.
					output = HttpResponse(
						json.dumps({ 'success' : {
							'message' : '''
							La prestation « {0} » a été reliée avec succès au dossier {1}.
							'''.format(o_prest, o_doss),
							'modal' : 'ajout_prest',
							'redirect' : reverse('cons_doss', args = [o_doss.pk])
						}}),
						content_type = 'application/json'
					)

					# Je renseigne l'onglet actif après rechargement de la page.
					request.session['tab_doss'] = '#ong_prest'

					if 'relier-prestation' in request.session :
						del request.session['relier-prestation']

					# Je complète le fichier log.
					rempl_fich_log([
						request.user.pk,
						request.user,
						o_doss.pk,
						o_doss,
						'C',
						'Reliage d\'une prestation',
						t_inst_prest_doss[-1].pk
					])

			# Je traite le cas où je dois supprimer un dossier.
			if get_action == 'supprimer-dossier' :
				output = HttpResponse(suppr(reverse('suppr_doss', args = [o_doss.pk])))

			# Ajout d'un élément dans la fiche de vie
			if get_action == 'ajouter-fdv' :

				# Soumission du formulaire
				f_ajout_fdv = GererFicheVie(request.POST, request.FILES, k_doss = o_doss)

				if f_ajout_fdv.is_valid() :

					# Création d'une instance TFicheVie
					fdv = f_ajout_fdv.save()

					# Affichage d'un message de succès
					output = HttpResponse(
						json.dumps({ 'success' : {
							'message' : 'La fiche de vie du dossier {} a été mise à jour avec succès.'.format(
								fdv.id_doss
							),
							'redirect' : reverse('cons_doss', args = [fdv.id_doss.pk])
						}}),
						content_type = 'application/json'
					)

					# Définition de l'onglet actif après rechargement de la page
					request.session['tab_doss'] = '#ong_fdv'

					# Remplissage du fichier log
					rempl_fich_log([
						request.user.pk, request.user, fdv.id_doss.pk, fdv.id_doss, 'C', 'Fiche de vie', fdv.pk
					])

				else :

					# Affichage des erreurs
					output = HttpResponse(json.dumps(f_ajout_fdv.errors), content_type = 'application/json')

			# Suppression d'un élément composant la fiche de vie
			if get_action == 'supprimer-fdv' :
				if request.GET.get('fdv') :
					output = HttpResponse(suppr(reverse('suppr_fdv', args = [request.GET['fdv']])))

	return output

'''
Cette vue permet de traiter le formulaire d'ajout d'un financement.
request : Objet requête
'''
@verif_acc
def ajout_fin(request) :

	# Imports
	from app.forms.gestion_dossiers import GererFinancement
	from app.functions import ger_droits
	from app.functions import rempl_fich_log
	from app.models import TDossier
	from django.core.urlresolvers import reverse
	from django.http import HttpResponse
	import json

	output = HttpResponse()

	if request.method == 'POST' :

		# Je vérifie le droit d'écriture.
		try :
			o_doss_droit = TDossier.objects.get(num_doss = request.POST.get('GererFinancement-za_num_doss'))
		except :
			o_doss_droit = None
		if o_doss_droit :
			ger_droits(request.user, o_doss_droit, False)

		# Je soumets le formulaire.
		f_ajout_fin = GererFinancement(request.POST, request.FILES, prefix = 'GererFinancement')

		if f_ajout_fin.is_valid() :

			# Je créé la nouvelle instance TFinancement.
			o_nv_fin = f_ajout_fin.save(commit = False)
			o_nv_fin.save()

			# J'affiche le message de succès.
			output = HttpResponse(
				json.dumps({ 'success' : {
					'message' : '''
					L\'organisme financeur « {0} » a été ajouté avec succès au plan de financement du dossier {1}.
					'''.format(o_nv_fin.id_org_fin, o_nv_fin.id_doss),
					'redirect' : reverse('cons_doss', args = [o_nv_fin.id_doss.pk])
				}}),
				content_type = 'application/json'
			)

			# Je renseigne l'onglet actif après rechargement de la page.
			request.session['tab_doss'] = '#ong_fin'

			# Je complète le fichier log.
			rempl_fich_log([
				request.user.pk, request.user, o_nv_fin.id_doss.pk, o_nv_fin.id_doss, 'C', 'Financement', o_nv_fin.pk
			])

		else :

			# J'affiche les erreurs.
			t_err = {}
			for k, v in f_ajout_fin.errors.items() :
				t_err['GererFinancement-{0}'.format(k)] = v
			output = HttpResponse(json.dumps(t_err), content_type = 'application/json')

	return output

'''
Cette vue permet d'afficher la page de modification d'un financement ou de traiter le formulaire de mise à jour d'un
financement.
request : Objet requête
_f : Identifiant d'un financement
'''
@verif_acc
def modif_fin(request, _f) :

	# Imports
	from app.forms.gestion_dossiers import GererFinancement
	from app.functions import ger_droits
	from app.functions import init_f
	from app.functions import init_fm
	from app.functions import rempl_fich_log
	from app.models import TFinancement
	from django.core.urlresolvers import reverse
	from django.http import HttpResponse
	from django.shortcuts import get_object_or_404
	from django.shortcuts import render
	import json

	output = HttpResponse()

	# Je vérifie l'existence d'un objet TFinancement.
	o_fin = get_object_or_404(TFinancement, pk = _f)

	# Je vérifie le droit d'écriture.
	ger_droits(request.user, o_fin.id_doss, False)

	if request.method == 'GET' :

		# J'instancie un objet "formulaire".
		f_modif_fin = GererFinancement(instance = o_fin, prefix = 'GererFinancement')

		# Je déclare un tableau de fenêtres modales.
		t_fm = [
			init_fm('modif_fin', 'Modifier un financement')
		]

		# J'affiche le template.
		output = render(
			request,
			'./gestion_dossiers/modif_fin.html',
			{ 'f' : o_fin, 'f_modif_fin' : init_f(f_modif_fin), 't_fm' : t_fm, 'title' : 'Modifier un financement' }
		)

	else :

		# Je soumets le formulaire.
		f_modif_fin = GererFinancement(request.POST, request.FILES, instance = o_fin, prefix = 'GererFinancement')

		if f_modif_fin.is_valid() :

			# Je modifie l'instance TFinancement.
			o_fin_modif = f_modif_fin.save(commit = False)
			o_fin_modif.save()

			# J'affiche le message de succès.
			output = HttpResponse(
				json.dumps({ 'success' : {
					'message' : 'Le plan de financement du dossier {0} a été mis à jour avec succès.'.format(
						o_fin_modif.id_doss
					),
					'redirect' : reverse('cons_fin', args = [o_fin_modif.pk])
				}}),
				content_type = 'application/json'
			)

			# Je complète le fichier log.
			rempl_fich_log([
				request.user.pk,
				request.user,
				o_fin_modif.id_doss.pk,
				o_fin_modif.id_doss,
				'U',
				'Financement',
				o_fin_modif.pk
			])

		else :

			# J'affiche les erreurs.
			t_err = {}
			for k, v in f_modif_fin.errors.items() :
				t_err['GererFinancement-{0}'.format(k)] = v
			output = HttpResponse(json.dumps(t_err), content_type = 'application/json')

	return output

'''
Cette vue permet de traiter le formulaire de suppression d'un financement.
request : Objet requête
_f : Identifiant d'un financement
'''
@verif_acc
@csrf_exempt
def suppr_fin(request, _f) :

	# Imports
	from app.functions import ger_droits
	from app.functions import rempl_fich_log
	from app.models import TDemandeVersement
	from app.models import TFinancement
	from django.core.urlresolvers import reverse
	from django.http import HttpResponse
	from django.shortcuts import get_object_or_404
	import json

	output = HttpResponse()

	# Je vérifie l'existence d'un objet TFinancement.
	o_fin = get_object_or_404(TFinancement, pk = _f)

	if request.method == 'POST' :

		# Je vérifie le droit d'écriture.
		ger_droits(request.user, o_fin.id_doss, False)

		# Je vérifie si je peux exécuter ou non la suppression du financement.
		qs_ddv = TDemandeVersement.objects.filter(id_fin = o_fin)
		if len(qs_ddv) == 0 :

			# Je pointe vers l'objet TFinancement à supprimer.
			o_fin_suppr = o_fin

			# Je récupère l'identifiant du financement afin de le renseigner dans le fichier log.
			v_id_fin = o_fin.pk

			# Je supprime l'objet TFinancement.
			o_fin.delete()

			# J'affiche le message de succès.
			output = HttpResponse(
				json.dumps({ 'success' : {
					'message' : '''
					L'organisme financeur « {0} » a été supprimé avec succès au plan de financement du dossier {1}.
					'''.format(o_fin_suppr.id_org_fin, o_fin_suppr.id_doss),
					'redirect' : reverse('cons_doss', args = [o_fin_suppr.id_doss.pk])
				}}),
				content_type = 'application/json'
			)

			# Je renseigne l'onglet actif après rechargement de la page.
			request.session['tab_doss'] = '#ong_fin'

			# Je complète le fichier log.
			rempl_fich_log([
				request.user.pk,
				request.user,
				o_fin_suppr.id_doss.pk,
				o_fin_suppr.id_doss,
				'D',
				'Financement',
				v_id_fin
			])

		else :
			cle = 'Demande de versement'
			if len(qs_ddv) > 1 :
				cle = 'Demandes de versements'

			# Je prépare le message d'alerte.
			mess_html = '''
			Veuillez d'abord supprimer le/les élément(s) suivant(s) :
			<ul class="list-inline">
				<li>{0} : {1}</li>
			</ul>
			'''.format(cle, len(qs_ddv))

			# J'affiche le message d'alerte.
			output = HttpResponse(json.dumps([mess_html]), content_type = 'application/json')

	return output

'''
Cette vue permet d'afficher la page de consultation d'un financement.
request : Objet requête
_f : Identifiant d'un financement
'''
@verif_acc
@csrf_exempt
def cons_fin(request, _f) :

	# Imports
	from app.functions import dt_fr
	from app.functions import init_pg_cons
	from app.functions import ger_droits
	from app.functions import init_f
	from app.functions import init_fm
	from app.functions import obt_mont
	from app.functions import obt_pourc
	from app.functions import suppr
	from app.models import TFinancement
	from app.models import VFinancement
	from django.core.urlresolvers import reverse
	from django.http import HttpResponse
	from django.shortcuts import get_object_or_404
	from django.shortcuts import render

	output = HttpResponse()

	# Je vérifie l'existence d'un objet TFinancement.
	o_fin = get_object_or_404(TFinancement, pk = _f)

	if request.method == 'GET' :

		# Je vérifie le droit de lecture.
		ger_droits(request.user, o_fin.id_doss)

		o_suivi_fin = VFinancement.objects.get(id_fin = o_fin.pk)

		# Je définis si le montant du financement est en HT ou en TTC.
		ht_ou_ttc = 'HT'
		if o_fin.id_doss.est_ttc_doss == True :
			ht_ou_ttc = 'TTC'

		# Je prépare le volet de consultation du financement.
		t_attrs_fin = {
			'id_doss' : { 'label' : 'Numéro du dossier', 'value' : o_fin.id_doss },
			'id_org_fin' : { 'label' : 'Organisme financeur', 'value' : o_fin.id_org_fin },
			'num_arr_fin' : { 'label' : 'Numéro de l\'arrêté ou convention', 'value' : o_fin.num_arr_fin or '' },
			'mont_elig_fin' : {
				'label' : 'Montant {0} de l\'assiette éligible de la subvention (en €)'.format(ht_ou_ttc),
				'value' : obt_mont(o_fin.mont_elig_fin) or ''
			},
			'pourc_elig_fin' : {
				'label' : 'Pourcentage de l\'assiette éligible',
				'value' : obt_pourc(o_fin.pourc_elig_fin) or ''
			},
			'pourc_glob_fin' : {
				'label' : 'Pourcentage global',
				'value' : obt_pourc(o_suivi_fin.pourc_glob_fin)
			},
			'mont_part_fin' : {
				'label' : 'Montant {0} total de la participation (en €)'.format(ht_ou_ttc),
				'value' : obt_mont(o_fin.mont_part_fin)
			},
			'dt_deb_elig_fin' : {
				'label' : 'Date de début d\'éligibilité', 'value' : dt_fr(o_fin.dt_deb_elig_fin) or ''
			},
			'dt_fin_elig_fin' : {
				'label' : 'Date de fin d\'éligibilité', 'value' : dt_fr(o_suivi_fin.dt_fin_elig_fin) or ''
			},
			'duree_valid_fin' : {
				'label' : 'Durée de validité de l\'aide (en mois)', 'value' : o_fin.duree_valid_fin or 0
			},
			'duree_pror_fin' : {
				'label' : 'Durée de la prorogation (en mois)', 'value' : o_fin.duree_pror_fin or 0
			},
			'dt_lim_deb_oper_fin' : {
				'label' : 'Date limite du début de l\'opération', 'value' : dt_fr(o_fin.dt_lim_deb_oper_fin) or ''
			},
			'a_inf_fin' : {
				'label' : 'Avez-vous informé le partenaire financier du début de l\'opération ?',
				'value' : o_fin.a_inf_fin
			},
			'dt_lim_prem_ac_fin' : {
				'label' : 'Date limite du premier acompte', 'value' : dt_fr(o_fin.dt_lim_prem_ac_fin) or ''
			},
			'id_paiem_prem_ac' : {
				'label' : 'Premier acompte payé en fonction de', 'value' : o_fin.id_paiem_prem_ac or ''
			},
			'pourc_real_fin' : {
				'label' : 'Pourcentage de réalisation des travaux', 'value' : obt_pourc(o_fin.pourc_real_fin) or ''
			},
			'chem_pj_fin' : {
				'label' : 'Consulter l\'arrêté de subvention', 'value' : o_fin.chem_pj_fin, 'pdf' : True
			},
			'comm_fin' : { 'label' : 'Commentaire', 'value' : o_fin.comm_fin or '' }
		}

		# Je déclare un tableau de fenêtres modales.
		t_fm = [
			init_fm('suppr_fin', 'Supprimer un financement')
		]

		# J'affiche le template.
		output = render(
			request,
			'./gestion_dossiers/cons_fin.html',
			{
				'f' : o_fin,
				'forbidden' : ger_droits(request.user, o_fin.id_doss, False, False),
				't_attrs_fin' : init_pg_cons(t_attrs_fin),
				't_fm' : t_fm,
				'title' : 'Consulter un financement'
			}
		)

	else :
		if 'action' in request.GET :

			# Je traite le cas où je dois supprimer un financement.
			if request.GET['action'] == 'supprimer-financement' :
				output = HttpResponse(suppr(reverse('suppr_fin', args = [o_fin.pk])))

	return output

'''
Cette vue permet de traiter le formulaire d'ajout d'une prestation.
request : Objet requête
'''
@verif_acc
def ajout_prest(request) :

	# Imports
	from app.forms.gestion_dossiers import GererPrestation
	from app.functions import ger_droits
	from app.functions import rempl_fich_log
	from app.models import TDossier
	from app.models import TPrestationsDossier
	from django.core.urlresolvers import reverse
	from django.http import HttpResponse
	import json

	output = HttpResponse()

	if request.method == 'POST' :

		# Je vérifie le droit d'écriture.
		try :
			o_doss_droit = TDossier.objects.get(num_doss = request.POST.get('GererPrestation-za_num_doss'))
		except :
			o_doss_droit = None
		if o_doss_droit :
			ger_droits(request.user, o_doss_droit, False)

		# Je soumets le formulaire.
		f_ajout_prest = GererPrestation(request.POST, request.FILES, prefix = 'GererPrestation')

		if f_ajout_prest.is_valid() :

			# Je récupère les données du formulaire valide.
			cleaned_data = f_ajout_prest.cleaned_data
			v_num_doss = cleaned_data.get('za_num_doss')
			v_duree_prest = cleaned_data.get('zs_duree_prest')
			v_mont_prest = cleaned_data.get('zs_mont_prest')

			# Je créé la nouvelle instance TPrestation.
			o_nvelle_prest = f_ajout_prest.save(commit = False)
			o_nvelle_prest.save()

			# Je fais le lien avec la table TPrestationsDossier.
			o_nvelle_prest_doss = TPrestationsDossier.objects.create(
				id_doss = TDossier.objects.get(num_doss = v_num_doss),
				id_prest = o_nvelle_prest,
				duree_prest_doss = v_duree_prest,
				mont_prest_doss = v_mont_prest
			)

			# J'affiche le message de succès.
			output = HttpResponse(
				json.dumps({ 'success' : {
					'message' : '''
					La prestation « {0} » a été ajoutée avec succès au dossier {1}.
					'''.format(o_nvelle_prest, v_num_doss),
					'redirect' : reverse('cons_doss', args = [o_nvelle_prest_doss.id_doss.pk])
				}}),
				content_type = 'application/json'
			)

			# Je renseigne l'onglet actif après rechargement de la page.
			request.session['tab_doss'] = '#ong_prest'

			# Je complète le fichier log.
			rempl_fich_log([
				request.user.pk,
				request.user,
				o_nvelle_prest_doss.id_doss.pk,
				o_nvelle_prest_doss.id_doss,
				'C',
				'Ajout d\'une prestation',
				o_nvelle_prest_doss.pk
			])

		else :

			# J'affiche les erreurs.
			t_err = {}
			for k, v in f_ajout_prest.errors.items() :
				t_err['GererPrestation-{0}'.format(k)] = v
			output = HttpResponse(json.dumps(t_err), content_type = 'application/json')

	return output

'''
Cette vue permet d'afficher la page de modification d'une prestation ou de traiter le formulaire de mise à jour d'une
prestation.
request : Objet requête
_p : Identifiant d'une prestation
'''
@verif_acc
def modif_prest(request, _pd) :

	# Imports
	from app.forms.gestion_dossiers import GererPrestation
	from app.forms.gestion_dossiers import RedistribuerPrestation
	from app.functions import ger_droits
	from app.functions import init_f
	from app.functions import init_fm
	from app.functions import obt_mont
	from app.functions import rempl_fich_log
	from app.models import TPrestationsDossier
	from app.models import VSuiviDossier
	from app.models import VSuiviPrestationsDossier
	from django.core.urlresolvers import reverse
	from django.http import HttpResponse
	from django.shortcuts import get_object_or_404
	from django.shortcuts import render
	from django.template.context_processors import csrf
	import json

	output = HttpResponse()

	# Je vérifie l'existence d'un objet TFacture.
	o_prest = get_object_or_404(TPrestationsDossier, pk = _pd)

	# Je vérifie le droit d'écriture.
	ger_droits(request.user, o_prest.id_doss, False)

	if request.method == 'GET' :

		# J'instancie un objet "formulaire".
		f_modif_prest = GererPrestation(instance = o_prest.id_prest, prefix = 'GererPrestation')

		t_lg = []
		for index, dp in enumerate(TPrestationsDossier.objects.filter(id_prest = o_prest.id_prest)) :

			# Je pointe vers l'objet VSuiviPrestationsDossier.
			o_suivi_prest_doss = VSuiviPrestationsDossier.objects.get(pk = dp.pk)

			# Je pointe vers l'objet VSuiviDossier.
			o_suivi_doss = VSuiviDossier.objects.get(pk = dp.id_doss.pk)

			# J'instancie un objet "formulaire".
			f_red_prest = RedistribuerPrestation(
				instance = dp, prefix = 'RedistribuerPrestation{0}'.format(index)
			)

			# J'initialise le gabarit de chaque champ du formulaire.
			t_red_prest = init_f(f_red_prest)

			# J'initialise la somme des factures émises selon le mode de taxe du dossier.
			mont_fact_sum = o_suivi_prest_doss.mont_ht_fact_sum
			if o_suivi_prest_doss.id_doss.est_ttc_doss == True :
				mont_fact_sum = o_suivi_prest_doss.mont_ttc_fact_sum

			lg = '''
			<tr>
				<td class="b">{0}</td>
				<td>{1}</td>
				<td>{2}</td>
				<td>{3}</td>
				<td>{4}</td>
				<td>{5}</td>
			</tr>
			'''.format(
				dp.id_doss.num_doss,
				t_red_prest['mont_prest_doss'],
				obt_mont(o_suivi_prest_doss.mont_aven_sum),
				obt_mont(mont_fact_sum),
				obt_mont(o_suivi_doss.mont_rae),
				t_red_prest['duree_prest_doss']
			)

			# J'empile le tableau des lignes du tableau HTML des dossiers déjà reliés à la prestation que l'on
			# veut modifier.
			t_lg.append(lg)

		# Je définis le mode de taxe du dossier.
		ht_ou_ttc = 'HT'
		if o_prest.id_doss.est_ttc_doss == True :
			ht_ou_ttc = 'TTC'

		# Je déclare un tableau de fenêtres modales.
		t_fm = [
			init_fm('modif_prest', 'Modifier une prestation')
		]

		# J'affiche le template.
		output = render(
			request,
			'./gestion_dossiers/modif_prest.html',
			{
				'f_modif_prest' : init_f(f_modif_prest),
				'ht_ou_ttc' : ht_ou_ttc,
				'pd' : o_prest,
				't_fm' : t_fm,
				't_lg' : '\n'.join(t_lg),
				'title' : 'Modifier une prestation'
			}
		)

	else :

		# Je déclare le tableau des instances TPrestation et TprestationsDossier à modifier.
		t_inst = []

		# J'initialise le tableau des erreurs tous formulaires confondus.
		t_err = {}

		# Je soumets le formulaire.
		f_modif_prest = GererPrestation(
			request.POST, request.FILES, instance = o_prest.id_prest, prefix = 'GererPrestation'
		)

		if f_modif_prest.is_valid() :

			# J'empile le tableau des instances.
			t_inst.append(f_modif_prest.save(commit = False))

		else :

			# J'empile le tableau des erreurs.
			for k, v in f_modif_prest.errors.items() :
				t_err['GererPrestation-{0}'.format(k)] = v

		for index, dp in enumerate(TPrestationsDossier.objects.filter(id_prest = o_prest.id_prest)) :

			# Je soumets le formulaire.
			f_red_prest = RedistribuerPrestation(
				request.POST, instance = dp, prefix = 'RedistribuerPrestation{0}'.format(index)
			)

			if f_red_prest.is_valid() :

				# J'empile le tableau des instances.
				t_inst.append(f_red_prest.save(commit = False))

			else :

				# J'empile le tableau des erreurs.
				for k, v in f_red_prest.errors.items() :
					t_err['RedistribuerPrestation{0}-{1}'.format(index, k)] = v

		if len(t_err) > 0 :

			# J'affiche les erreurs.
			output = HttpResponse(json.dumps(t_err), content_type = 'application/json')

		else :

			# Je modifie chaque instance.
			for i in range(0, len(t_inst)) :
				t_inst[i].save()

			# J'affiche le message de succès.
			output = HttpResponse(
				json.dumps({ 'success' : {
					'message' : 'La prestation « {0} » a été mise à jour avec succès.'.format(o_prest.id_prest),
					'redirect' : reverse('cons_prest', args = [o_prest.pk])
				}}),
				content_type = 'application/json'
			)

			# Je complète le fichier log.
			rempl_fich_log([
				request.user.pk, request.user, o_prest.id_doss.pk, o_prest.id_doss, 'U', 'Prestation', o_prest.pk
			])

	return output

'''
Cette vue permet de traiter le formulaire de suppression d'une prestation.
request : Objet requête
_pd : Identifiant d'un couple prestation/dossier
'''
@verif_acc
@csrf_exempt
def suppr_prest(request, _pd) :

	# Imports
	from app.functions import ger_droits
	from app.functions import rempl_fich_log
	from app.models import TAvenant
	from app.models import TFacture
	from app.models import TPrestationsDossier
	from django.core.urlresolvers import reverse
	from django.http import HttpResponse
	from django.shortcuts import get_object_or_404
	import json

	output = HttpResponse()

	# Je vérifie l'existence d'un objet TPrestationsDossier.
	o_prest_doss = get_object_or_404(TPrestationsDossier, pk = _pd)

	if request.method == 'POST' :

		# Je vérifie le droit d'écriture.
		ger_droits(request.user, o_prest_doss.id_doss, False)

		# J'initialise un tableau de jeu de données.
		t_qs = [
			['Avenant', TAvenant.objects.filter(id_doss = o_prest_doss.id_doss, id_prest = o_prest_doss.id_prest)],
			['Facture', TFacture.objects.filter(id_doss = o_prest_doss.id_doss, id_prest = o_prest_doss.id_prest)]
		]

		# Je vérifie si je peux exécuter ou non la suppression du couple prestation/dossier.
		peut_suppr = True
		t_elem_a_suppr = []
		for i in range(0, len(t_qs)) :
			if len(t_qs[i][1]) > 0 :
				peut_suppr = False
				cle = t_qs[i][0]
				if len(t_qs[i][1]) > 1 :
					split = cle.split(' ')
					for j in range(0, len(split)) :
						split[j] += 's'
					cle = ' '.join(split)
				t_elem_a_suppr.append('{0} : {1}'.format(cle, len(t_qs[i][1])))

		if peut_suppr == True :

			# Je pointe vers l'objet TPrestationsDossier à supprimer.
			o_prest_doss_suppr = o_prest_doss

			# Je récupère l'identifiant du couple prestation/dossier afin de le renseigner dans le fichier log.
			v_id_prest_doss = o_prest_doss.pk

			# Je supprime l'objet TPrestationsDossier.
			o_prest_doss.delete()

			# Je supprime l'objet TPrestation s'il n'est lié à aucun autre objet TPrestationsDossier.
			if len(TPrestationsDossier.objects.filter(id_prest = o_prest_doss_suppr.id_prest)) == 0 :
				o_prest_doss.id_prest.delete()

			# J'affiche le message de succès.
			output = HttpResponse(
				json.dumps({ 'success' : {
					'message' : 'La prestation « {0} » a été supprimée avec succès.'.format(
						o_prest_doss_suppr.id_prest
					),
					'redirect' : reverse('cons_doss', args = [o_prest_doss_suppr.id_doss.pk])
				}}),
				content_type = 'application/json'
			)

			# Je renseigne l'onglet actif après rechargement de la page.
			request.session['tab_doss'] = '#ong_prest'

			# Je complète le fichier log.
			rempl_fich_log([
				request.user.pk,
				request.user,
				o_prest_doss_suppr.id_doss.pk,
				o_prest_doss_suppr.id_doss,
				'D',
				'Prestation',
				v_id_prest_doss
			])

		else :

			# Je prépare le message d'alerte.
			mess_html = 'Veuillez d\'abord supprimer le/les élément(s) suivant(s) :'
			mess_html += '<ul class="list-inline">'
			for i in range(0, len(t_elem_a_suppr)) :
				mess_html += '<li>{0}</li>'.format(t_elem_a_suppr[i])
			mess_html += '</ul>'

			# J'affiche le message d'alerte.
			output = HttpResponse(json.dumps([mess_html]), content_type = 'application/json')

	return output

'''
Cette vue permet d'afficher la page de consultation d'une prestation.
request : Objet requête
_p : Identifiant d'une prestation
'''
@verif_acc
@csrf_exempt
def cons_prest(request, _pd) :

	# Imports
	from app.forms.gestion_dossiers import GererOrdreService
	from app.functions import dt_fr
	from app.functions import gen_f_ajout_aven
	from app.functions import ger_droits
	from app.functions import init_fm
	from app.functions import init_pg_cons
	from app.functions import obt_mont
	from app.functions import suppr
	from app.models import TAvenant
	from app.models import TOrdreService
	from app.models import TPrestationsDossier
	from app.models import VPrestation
	from app.models import VSuiviPrestationsDossier
	from django.core.urlresolvers import reverse
	from django.http import HttpResponse
	from django.shortcuts import get_object_or_404
	from django.shortcuts import render
	import json

	output = HttpResponse()

	# Je vérifie l'existence d'un objet TPrestationsDossier.
	o_prest_doss = get_object_or_404(TPrestationsDossier, pk = _pd)

	if request.method == 'GET' :

		# Je vérifie le droit de lecture.
		ger_droits(request.user, o_prest_doss.id_doss)

		# Je désigne l'onglet actif du navigateur à onglets relatif à une prestation.
		if 'tab_prest' not in request.session :
			request.session['tab_prest'] = '#ong_prest'

		# Je définis si le montant de la prestation est en HT ou en TTC.
		ht_ou_ttc = 'HT'
		if o_prest_doss.id_doss.est_ttc_doss == True :
			ht_ou_ttc = 'TTC'

		# Je pointe vers l'objet VSuiviPrestationsDossier.
		o_suivi_prest_doss = VSuiviPrestationsDossier.objects.get(pk = o_prest_doss.pk)

		# J'initialise la valeur attributaire liée à la somme des factures émises.
		v_mont_fact_sum = o_suivi_prest_doss.mont_ht_fact_sum
		if ht_ou_ttc == 'TTC' :
			v_mont_fact_sum = o_suivi_prest_doss.mont_ttc_fact_sum

		# Je pointe vers l'objet VPrestation.
		o_suivi_prest = VPrestation.objects.get(pk = o_prest_doss.id_prest.pk)

		# Je prépare le volet de consultation de la prestation.
		t_attrs_prest_doss = {
			'id_org_prest' : { 'label': 'Prestataire', 'value' : o_suivi_prest.get_instance().id_org_prest },
			'int_prest' : { 'label': 'Intitulé de la prestation', 'value' : o_suivi_prest.get_instance().int_prest },
			'ref_prest' : { 'label': 'Référence de la prestation', 'value' : o_suivi_prest.get_instance().ref_prest },
			'mont_prest' : {
				'label': 'Montant {0} total de la prestation (en €)'.format(ht_ou_ttc),
				'value' : obt_mont(o_suivi_prest.mont_prest)
			},
			'dt_notif_prest' : {
				'label' : 'Date de notification de la prestation',
				'value' : dt_fr(o_suivi_prest.get_instance().dt_notif_prest) or ''
			},
			'dt_fin_prest' : {
				'label' : 'Date de fin de la prestation',
				'value' : dt_fr(o_suivi_prest.get_instance().dt_fin_prest) or ''
			},
			'id_nat_prest' : {
				'label' : 'Nature de la prestation', 'value' : o_suivi_prest.get_instance().id_nat_prest
			},
			'chem_pj_prest' : {
				'label' : 'Consulter le contrat de prestation',
				'value' : o_suivi_prest.get_instance().chem_pj_prest,
				'pdf' : True
			},
			'comm_prest' : { 'label' : 'Commentaire', 'value' : o_suivi_prest.get_instance().comm_prest or '' },
			'duree_prest_doss' : {
				'label' : 'Durée de la prestation (en nombre de jours ouvrés)',
				'value' : o_prest_doss.duree_prest_doss
			},
			'duree_w_os_sum' : {
				'label' : 'Durée travaillée (en nombre de jours ouvrés)',
				'value' : o_suivi_prest_doss.duree_w_os_sum
			},
			'duree_w_rest_os_sum' : {
				'label' : 'Durée restante à travailler (en nombre de jours ouvrés)',
				'value' : o_suivi_prest_doss.duree_w_rest_os_sum
			},
			'mont_prest_doss' : {
				'label' : 'Montant {0} de la prestation (en €)'.format(ht_ou_ttc),
				'value' : obt_mont(o_prest_doss.mont_prest_doss)
			},
			'nb_aven' : { 'label' : 'Nombre d\'avenants', 'value' : str(o_suivi_prest_doss.nb_aven) },
			'mont_aven_sum' : {
				'label' : 'Somme {0} des avenants (en €)'.format(ht_ou_ttc),
				'value' : obt_mont(o_suivi_prest_doss.mont_aven_sum)
			},
			'mont_fact_sum' : {
				'label' : 'Somme {0} des factures émises (en €)'.format(ht_ou_ttc),
				'value' : obt_mont(v_mont_fact_sum)
			},
			'mont_raf' : {
				'label' : 'Reste à facturer {0} (en €)'.format(ht_ou_ttc),
				'value' : obt_mont(o_suivi_prest_doss.mont_raf),
				'help-text' : '''
				= montant {ht_ou_ttc} de la prestation + somme {ht_ou_ttc} des avenants - somme {ht_ou_ttc} des
				factures émises
				'''.format(ht_ou_ttc = ht_ou_ttc)
			}
		}

		# J'initialise le tableau des autres dossiers reliés à la prestation.
		t_doss = [{
			'id_doss' : d.id_doss,
			'mont_prest_doss' : obt_mont(d.mont_prest_doss),
			'mont_aven_sum' : obt_mont(d.mont_aven_sum),
			'pk' : d.pk
		} for d in VSuiviPrestationsDossier.objects.filter(id_prest = o_prest_doss.id_prest).exclude(
			pk = o_prest_doss.pk
		).order_by('id_doss')]

		# J'initialise le tableau des avenants du couple prestation/dossier.
		t_aven = [{
			'num_aven' : index + 1,
			'int_aven' : a.int_aven,
			'dt_aven' : dt_fr(a.dt_aven) or '-',
			'mont_aven' : obt_mont(a.mont_aven) or 0,
			'pk' : a.pk
		} for index, a in enumerate(
			TAvenant.objects.filter(id_doss = o_prest_doss.id_doss, id_prest = o_prest_doss.id_prest)
		)]

		# J'initialise le tableau des ordres de service du couple prestation/dossier.
		t_os = [{
			'pk' : os.pk,
			'comm_os' : os.comm_os,
			'd_emiss_os' : dt_fr(os.d_emiss_os),
			'duree_w_os' : os.duree_w_os,
			'num_os' : os.num_os,
			'obj_os' : os.obj_os,
			'id_type_os' : os.id_type_os
		} for os in TOrdreService.objects.filter(id_doss = o_prest_doss.id_doss, id_prest = o_prest_doss.id_prest)]

		# Je déclare un tableau qui stocke le contenu de certaines fenêtres modales.
		t_cont_fm = {
			'ajout_aven' : gen_f_ajout_aven(request, o_prest_doss, reverse('ajout_aven'))
		}

		# Je déclare un tableau de fenêtres modales.
		t_fm = [
			init_fm('ajout_aven', 'Ajouter un avenant', t_cont_fm['ajout_aven']),
			init_fm(
				'geros',
				'Ajouter un ordre de service',
				GererOrdreService(k_pd = o_prest_doss, prefix = 'GererOrdreService').get_form(rq = request)
			),
			init_fm('suppr_os', 'Supprimer un ordre de service'),
			init_fm('suppr_prest', 'Supprimer une prestation')
		]

		# J'affiche le template.
		output = render(
			request,
			'./gestion_dossiers/cons_prest.html',
			{
				'pd' : o_prest_doss,
				'duree_w_os_sum' : o_suivi_prest_doss.duree_w_os_sum,
				'duree_w_rest_os_sum' : o_suivi_prest_doss.duree_w_rest_os_sum,
				'forbidden' : ger_droits(request.user, o_prest_doss.id_doss, False, False),
				'ht_ou_ttc' : ht_ou_ttc,
				't_attrs_prest_doss' : init_pg_cons(t_attrs_prest_doss),
				't_aven' : t_aven,
				't_doss' : t_doss,
				't_fm' : t_fm,
				't_os' : t_os,
				'title' : 'Consulter une prestation'
			}
		)

		# Je supprime la variable de session liée au navigateur à onglets relatif à une prestation.
		del request.session['tab_prest']

	else :
		if 'action' in request.GET :

			# Suppression d'un ordre de service
			if request.GET['action'] == 'supprimer-os' :
				if request.GET.get('os') :
					output = HttpResponse(suppr(reverse('suppr_os', args = [request.GET['os']])))

			# Je traite le cas où je dois supprimer une prestation.
			if request.GET['action'] == 'supprimer-prestation' :

				# Je prépare un message d'avertissement au cas où la prestation est reliée à plusieurs dossiers.
				avert = ''
				if len(TPrestationsDossier.objects.filter(id_prest = o_prest_doss.id_prest)) > 1 :
					avert = '''
					La suppression de la prestation n\'entraînera pas la suppression totale de celle-ci car elle est
					reliée au minimum à un autre dossier.
					'''

				output = HttpResponse(suppr(reverse('suppr_prest', args = [o_prest_doss.pk]), avert))

	return output

'''
Cette vue permet de traiter le formulaire d'ajout d'un avenant.
request : Objet requête
_r : Paramètres de redirection
'''
@verif_acc
def ajout_aven(request, _r) :

	# Imports
	from app.forms.gestion_dossiers import GererAvenant
	from app.functions import ger_droits
	from app.functions import rempl_fich_log
	from app.models import TDossier
	from app.models import TPrestationsDossier
	from django.core.urlresolvers import reverse
	from django.http import HttpResponse
	import json

	output = HttpResponse()

	if request.method == 'POST' :

		# Je vérifie le droit d'écriture.
		try :
			o_doss_droit = TDossier.objects.get(num_doss = request.POST.get('GererAvenant-za_num_doss'))
		except :
			o_doss_droit = None
		if o_doss_droit :
			ger_droits(request.user, o_doss_droit, False)

		# Je soumets le formulaire.
		f_ajout_aven = GererAvenant(request.POST, request.FILES, prefix = 'GererAvenant')

		# Je rajoute des choix valides pour la liste déroulante des prestations (prévention d'erreurs).
		post_prest = request.POST.get('GererAvenant-zl_prest')
		if post_prest :
			f_ajout_aven.fields['zl_prest'].choices = [(post_prest, post_prest)]

		if f_ajout_aven.is_valid() :

			# Je récupère les données du formulaire valide.
			cleaned_data = f_ajout_aven.cleaned_data
			v_num_doss = cleaned_data.get('za_num_doss')
			v_prest = cleaned_data.get('zl_prest')

			# Je pointe vers l'objet TPrestationsDossier.
			o_prest_doss = TPrestationsDossier.objects.get(id_doss__num_doss = v_num_doss, id_prest = v_prest)

			# Je créé la nouvelle instance TAvenant.
			o_nv_aven = f_ajout_aven.save(commit = False)
			o_nv_aven.num_aven = o_prest_doss.seq_aven_prest_doss
			o_nv_aven.save()

			# Je mets à jour l'objet TPrestationsDossier.
			o_prest_doss.seq_aven_prest_doss += 1
			o_prest_doss.save()

			if _r == 'cons_prest' :
				url = reverse(
					'cons_prest',
					args = [TPrestationsDossier.objects.get(
						id_doss = o_nv_aven.id_doss, id_prest = o_nv_aven.id_prest
					).pk]
				)
				tab = 'tab_prest'
				ong = '#ong_aven'
			if _r == 'cons_doss' :
				url = reverse(
					'cons_doss',
					args = [o_nv_aven.id_doss.pk]
				)
				tab = 'tab_doss'
				ong = '#ong_prest'

			# J'affiche le message de succès.
			output = HttpResponse(
				json.dumps({ 'success' : {
					'message' : 'L\'avenant a été ajouté avec succès à la prestation « {0} » du dossier {1}.'.format(
						o_nv_aven.id_prest, o_nv_aven.id_doss
					),
					'redirect' : url
				}}),
				content_type = 'application/json'
			)

			# Je renseigne l'onglet actif après rechargement de la page.
			request.session[tab] = ong

			# Je complète le fichier log.
			rempl_fich_log([
				request.user.pk, request.user, o_nv_aven.id_doss.pk, o_nv_aven.id_doss, 'C', 'Avenant', o_nv_aven.pk
			])

		else :

			# J'affiche les erreurs.
			t_err = {}
			for k, v in f_ajout_aven.errors.items() :
				t_err['GererAvenant-{0}'.format(k)] = v
			output = HttpResponse(json.dumps(t_err), content_type = 'application/json')

	return output

'''
Cette vue permet d'afficher la page de modification d'un avenant ou de traiter le formulaire de mise à jour d'un
avenant
request : Objet requête
_a : Identifiant d'un avenant
'''
@verif_acc
def modif_aven(request, _a) :

	# Imports
	from app.forms.gestion_dossiers import GererAvenant
	from app.functions import ger_droits
	from app.functions import init_f
	from app.functions import init_fm
	from app.functions import rempl_fich_log
	from app.models import TAvenant
	from django.core.urlresolvers import reverse
	from django.http import HttpResponse
	from django.shortcuts import get_object_or_404
	from django.shortcuts import render
	import json

	output = HttpResponse()

	# Je vérifie l'existence d'un objet TAvenant.
	o_aven = get_object_or_404(TAvenant, pk = _a)

	# Je vérifie le droit d'écriture.
	ger_droits(request.user, o_aven.id_doss, False)

	if request.method == 'GET' :

		# J'instancie un objet "formulaire".
		f_modif_aven = GererAvenant(instance = o_aven, prefix = 'GererAvenant')

		# Je déclare un tableau de fenêtres modales.
		t_fm = [
			init_fm('modif_aven', 'Modifier un avenant')
		]

		# J'affiche le template.
		output = render(
			request,
			'./gestion_dossiers/modif_aven.html',
			{ 'a' : o_aven, 'f_modif_aven' : init_f(f_modif_aven), 't_fm' : t_fm, 'title' : 'Modifier un avenant' }
		)

	else :

		# Je soumets le formulaire.
		f_modif_aven = GererAvenant(request.POST, instance = o_aven, prefix = 'GererAvenant')

		if f_modif_aven.is_valid() :

			# Je modifie l'instance TAvenant.
			o_aven_modif = f_modif_aven.save(commit = False)
			o_aven_modif.save()

			# J'affiche le message de succès.
			output = HttpResponse(
				json.dumps({ 'success' : {
					'message' : 'L\'avenant a été mis à jour avec succès.',
					'redirect' : reverse('cons_aven', args = [o_aven_modif.pk])
				}}),
				content_type = 'application/json'
			)

			# Je complète le fichier log.
			rempl_fich_log([
				request.user.pk,
				request.user,
				o_aven_modif.id_doss.pk,
				o_aven_modif.id_doss,
				'U',
				'Avenant',
				o_aven_modif.pk
			])

		else :

			# J'affiche les erreurs.
			t_err = {}
			for k, v in f_modif_aven.errors.items() :
				t_err['GererAvenant-{0}'.format(k)] = v
			output = HttpResponse(json.dumps(t_err), content_type = 'application/json')

	return output

'''
Cette vue permet de traiter le formulaire de suppression d'un avenant.
request : Objet requête
_a : Identifiant d'un avenant
'''
@verif_acc
@csrf_exempt
def suppr_aven(request, _a) :

	# Imports
	from app.functions import ger_droits
	from app.functions import rempl_fich_log
	from app.models import TAvenant
	from app.models import TPrestationsDossier
	from app.models import VSuiviPrestationsDossier
	from django.core.urlresolvers import reverse
	from django.db.models import F
	from django.http import HttpResponse
	from django.shortcuts import get_object_or_404
	import json

	output = HttpResponse()

	# Je vérifie l'existence d'un objet TAvenant.
	o_aven = get_object_or_404(TAvenant, pk = _a)

	if request.method == 'POST' :

		# Je vérifie le droit d'écriture.
		ger_droits(request.user, o_aven.id_doss, False)

		# Je pointe vers l'objet VSuiviPrestationsDossier.
		o_suivi_prest_doss = VSuiviPrestationsDossier.objects.get(id_doss = o_aven.id_doss, id_prest = o_aven.id_prest)

		# Je mets à jour le montant de la prestation si incohérence il y a.
		if o_aven.mont_aven and o_suivi_prest_doss.mont_raf < o_aven.mont_aven :
			TPrestationsDossier.objects.filter(pk = o_suivi_prest_doss.pk).update(
				mont_prest_doss = F('mont_prest_doss') + o_aven.mont_aven
			)

		# Je pointe vers l'objet TAvenant à supprimer.
		o_aven_suppr = o_aven

		# Je récupère l'identifiant de l'avenant afin de le renseigner dans le fichier log.
		v_id_aven = o_aven.pk

		# Je supprime l'objet TAvenant.
		o_aven.delete()

		# J'affiche le message de succès.
		output = HttpResponse(
			json.dumps({ 'success' : {
				'message' : 'L\'avenant a été supprimé avec succès.',
				'redirect' : reverse('cons_prest', args = [o_suivi_prest_doss.pk])
			}}),
			content_type = 'application/json'
		)

		# Je renseigne l'onglet actif après rechargement de la page.
		request.session['tab_prest'] = '#ong_aven'

		# Je complète le fichier log.
		rempl_fich_log([
			request.user.pk, request.user, o_aven_suppr.id_doss.pk, o_aven_suppr.id_doss, 'D', 'Avenant', v_id_aven
		])

	return output

'''
Cette vue permet d'afficher la page de consultation d'un avenant.
request : Objet requête
_a : Identifiant d'un avenant
'''
@verif_acc
@csrf_exempt
def cons_aven(request, _a) :

	# Imports
	from app.functions import dt_fr
	from app.functions import ger_droits
	from app.functions import init_fm
	from app.functions import init_pg_cons
	from app.functions import obt_mont
	from app.functions import suppr
	from app.models import TAvenant
	from app.models import TPrestationsDossier
	from app.models import VSuiviPrestationsDossier
	from django.http import HttpResponse
	from django.core.urlresolvers import reverse
	from django.shortcuts import get_object_or_404
	from django.shortcuts import render

	output = HttpResponse()

	# Je vérifie l'existence d'un objet TAvenant.
	o_aven = get_object_or_404(TAvenant, pk = _a)

	# Je pointe vers l'objet TPrestationsDossier.
	o_prest_doss = TPrestationsDossier.objects.get(id_doss = o_aven.id_doss, id_prest = o_aven.id_prest)

	if request.method == 'GET' :

		# Je vérifie le droit de lecture.
		ger_droits(request.user, o_aven.id_doss)

		# Je définis si le montant de l'avenant est en HT ou en TTC.
		ht_ou_ttc = 'HT'
		if o_aven.id_doss.est_ttc_doss == True :
			ht_ou_ttc = 'TTC'

		# Je prépare le volet de consultation de l'avenant.
		t_attrs_aven = {
			'id_doss' : { 'label': 'Numéro du dossier', 'value' : o_aven.id_doss },
			'id_prest' : { 'label' : 'Prestation', 'value' : o_aven.id_prest },
			'int_aven' : { 'label' : 'Intitulé de l\'avenant', 'value' : o_aven.int_aven },
			'dt_aven' : { 'label' : 'Date de fin de l\'avenant', 'value' : dt_fr(o_aven.dt_aven) or '' },
			'mont_aven' : {
				'label' : 'Montant {0} de l\'avenant (en €)'.format(ht_ou_ttc),
				'value' : obt_mont(o_aven.mont_aven) or 0
			},
			'chem_pj_aven' : {
				'label' : 'Consulter le fichier scanné de l\'avenant', 'value' : o_aven.chem_pj_aven, 'pdf' : True
			},
			'comm_aven' : { 'label' : 'Commentaire', 'value' : o_aven.comm_aven or '' }
		}

		# Je déclare un tableau de fenêtres modales.
		t_fm = [
			init_fm('suppr_aven', 'Supprimer un avenant')
		]

		# J'affiche le template.
		output = render(
			request,
			'./gestion_dossiers/cons_aven.html',
			{
				'a' : o_aven,
				'forbidden' : ger_droits(request.user, o_aven.id_doss, False, False),
				'ht_ou_ttc' : ht_ou_ttc,
				'pd' : o_prest_doss,
				't_attrs_aven' : init_pg_cons(t_attrs_aven),
				't_fm' : t_fm,
				'title' : 'Consulter un avenant'
			}
		)

	else :
		if 'action' in request.GET :

			# Je traite le cas où je dois supprimer un avenant.
			if request.GET['action'] == 'supprimer-avenant' :

				# Je définis le mode de taxe du dossier.
				ht_ou_ttc = 'HT'
				if o_aven.id_doss.est_ttc_doss == True :
					ht_ou_ttc = 'TTC'

				# Je prépare un message d'avertissement au cas où l'avenant couvre la facturation.
				avert = ''
				if o_aven.mont_aven and VSuiviPrestationsDossier.objects.get(
					pk = o_prest_doss.pk
				).mont_raf < o_aven.mont_aven :
					avert = '''
					Afin d'éviter une incohérence dans les chiffres, la suppression de l'avenant entraînera une
					modification du montant {ht_ou_ttc} de la prestation « {prest} », causée par la facturation de
					celle-ci. Si vous supprimez cet avenant, alors le montant de celui-ci sera transféré vers la
					prestation « {prest} ».
					'''.format(ht_ou_ttc = ht_ou_ttc, prest = o_prest_doss.id_prest)

				output = HttpResponse(suppr(reverse('suppr_aven', args = [o_aven.pk]), avert))

	return output

'''
Cette vue permet de traiter le formulaire d'ajout d'une facture.
request : Objet requête
'''
@verif_acc
def ajout_fact(request) :

	# Imports
	from app.forms.gestion_dossiers import GererFacture
	from app.functions import ger_droits
	from app.functions import rempl_fich_log
	from app.models import TDossier
	from app.models import TPrestation
	from django.core.urlresolvers import reverse
	from django.http import HttpResponse
	import json

	output = HttpResponse()

	if request.method == 'POST' :

		# Je vérifie le droit d'écriture.
		try :
			o_doss_droit = TDossier.objects.get(num_doss = request.POST.get('GererFacture-za_num_doss'))
		except :
			o_doss_droit = None
		if o_doss_droit :
			ger_droits(request.user, o_doss_droit, False)

		# Je soumets le formulaire.
		f_ajout_fact = GererFacture(request.POST, request.FILES, prefix = 'GererFacture')

		# Je rajoute des choix valides pour la liste déroulante des prestations (prévention d'erreurs).
		post_prest = request.POST.get('GererFacture-zl_prest')
		if post_prest :
			f_ajout_fact.fields['zl_prest'].choices = [(p.pk, p) for p in TPrestation.objects.filter(pk = post_prest)]

		if f_ajout_fact.is_valid() :

			# Je créé la nouvelle instance TFacture.
			o_nvelle_fact = f_ajout_fact.save(commit = False)
			o_nvelle_fact.save()

			# J'affiche le message de succès.
			output = HttpResponse(
				json.dumps({ 'success' : {
					'message' : '''
					La facture N°{0} a été ajoutée avec succès à la prestation « {1} » du dossier {2}.
					'''.format(o_nvelle_fact.num_fact, o_nvelle_fact.id_prest, o_nvelle_fact.id_doss),
					'redirect' : reverse('cons_doss', args = [o_nvelle_fact.id_doss.pk])
				}}),
				content_type = 'application/json'
			)

			# Je renseigne l'onglet actif après rechargement de la page.
			request.session['tab_doss'] = '#ong_fact'

			# Je complète le fichier log.
			rempl_fich_log([
				request.user.pk,
				request.user,
				o_nvelle_fact.id_doss.pk,
				o_nvelle_fact.id_doss,
				'C',
				'Facture',
				o_nvelle_fact.pk
			])

		else :

			# J'affiche les erreurs.
			t_err = {}
			for k, v in f_ajout_fact.errors.items() :
				t_err['GererFacture-{0}'.format(k)] = v
			output = HttpResponse(json.dumps(t_err), content_type = 'application/json')

	return output

'''
Cette vue permet d'afficher la page de modification d'une facture ou de traiter le formulaire de mise à jour d'une
facture.
request : Objet requête
_f : Identifiant d'une facture
'''
@verif_acc
def modif_fact(request, _f) :

	# Imports
	from app.forms.gestion_dossiers import GererFacture
	from app.functions import ger_droits
	from app.functions import init_f
	from app.functions import init_fm
	from app.functions import rempl_fich_log
	from app.models import TFacture
	from django.core.urlresolvers import reverse
	from django.http import HttpResponse
	from django.shortcuts import get_object_or_404
	from django.shortcuts import render
	import json

	output = HttpResponse()

	# Je vérifie l'existence d'un objet TFacture.
	o_fact = get_object_or_404(TFacture, pk = _f)

	# Je vérifie le droit d'écriture.
	ger_droits(request.user, o_fact.id_doss, False)

	if request.method == 'GET' :

		# J'instancie un objet "formulaire".
		f_modif_fact = GererFacture(instance = o_fact, prefix = 'GererFacture')

		# Je déclare un tableau de fenêtres modales.
		t_fm = [
			init_fm('modif_fact', 'Modifier une facture')
		]

		# J'affiche le template.
		output = render(
			request,
			'./gestion_dossiers/modif_fact.html',
			{ 'f' : o_fact, 'f_modif_fact' : init_f(f_modif_fact), 't_fm' : t_fm, 'title' : 'Modifier une facture' }
		)

	else :

		# Je soumets le formulaire.
		f_modif_fact = GererFacture(request.POST, request.FILES, instance = o_fact, prefix = 'GererFacture')

		if f_modif_fact.is_valid() :

			# Je modifie l'instance TFacture.
			o_fact_modif = f_modif_fact.save(commit = False)
			o_fact_modif.save()

			# J'affiche le message de succès.
			output = HttpResponse(
				json.dumps({ 'success' : {
					'message' : 'La facture N°{0} a été mise à jour avec succès.'.format(o_fact_modif.num_fact),
					'redirect' : reverse('cons_fact', args = [o_fact_modif.pk])
				}}),
				content_type = 'application/json'
			)

			# Je complète le fichier log.
			rempl_fich_log([
				request.user.pk,
				request.user,
				o_fact_modif.id_doss.pk,
				o_fact_modif.id_doss,
				'U',
				'Facture',
				o_fact_modif.pk
			])

		else :

			# J'affiche les erreurs.
			t_err = {}
			for k, v in f_modif_fact.errors.items() :
				t_err['GererFacture-{0}'.format(k)] = v
			output = HttpResponse(json.dumps(t_err), content_type = 'application/json')

	return output

'''
Cette vue permet de traiter le formulaire de suppression d'une facture.
request : Objet requête
_f : Identifiant d'une facture
'''
@verif_acc
@csrf_exempt
def suppr_fact(request, _f) :

	# Imports
	from app.functions import ger_droits
	from app.functions import rempl_fich_log
	from app.models import TFacture
	from app.models import TFacturesDemandeVersement
	from django.core.urlresolvers import reverse
	from django.http import HttpResponse
	from django.shortcuts import get_object_or_404
	import json

	output = HttpResponse()

	# Je vérifie l'existence d'un objet TFacture.
	o_fact = get_object_or_404(TFacture, pk = _f)

	if request.method == 'POST' :

		# Je vérifie le droit d'écriture.
		ger_droits(request.user, o_fact.id_doss, False)

		# Je vérifie si je peux exécuter ou non la suppression de la facture.
		qs_fddv = TFacturesDemandeVersement.objects.filter(id_fact = o_fact)
		if len(qs_fddv) == 0 :

			# Je pointe vers l'objet TFacture à supprimer.
			o_fact_suppr = o_fact

			# Je récupère l'identifiant de la facture afin de le renseigner dans le fichier log.
			v_id_fact = o_fact.pk

			# Je supprime l'objet TFacture.
			o_fact.delete()

			# J'affiche le message de succès.
			output = HttpResponse(
				json.dumps({ 'success' : {
					'message' : 'La facture N°{0} a été supprimée avec succès.'.format(o_fact_suppr),
					'redirect' : reverse('cons_doss', args = [o_fact_suppr.id_doss.pk])
				}}),
				content_type = 'application/json'
			)

			# Je renseigne l'onglet actif après rechargement de la page.
			request.session['tab_doss'] = '#ong_fact'

			# Je complète le fichier log.
			rempl_fich_log([
				request.user.pk,
				request.user,
				o_fact_suppr.id_doss.pk,
				o_fact_suppr.id_doss,
				'D',
				'Facture',
				v_id_fact
			])

		else :
			cle = 'Demande de versement'
			if len(qs_fddv) > 1 :
				cle = 'Demandes de versements'

			# Je prépare le message d'alerte.
			mess_html = '''
			Veuillez d'abord modifier ou supprimer le/les élément(s) suivant(s) :
			<ul class="list-inline">
				<li>{0} : {1}</li>
			</ul>
			'''.format(cle, len(qs_fddv))

			# J'affiche le message d'alerte.
			output = HttpResponse(json.dumps([mess_html]), content_type = 'application/json')

	return output

'''
Cette vue permet d'afficher la page de consultation d'une facture.
request : Objet requête
_f : Identifiant d'une facture
'''
@verif_acc
@csrf_exempt
def cons_fact(request, _f) :

	# Imports
	from app.functions import dt_fr
	from app.functions import ger_droits
	from app.functions import init_fm
	from app.functions import init_pg_cons
	from app.functions import obt_mont
	from app.functions import suppr
	from app.models import TFacture
	from django.core.urlresolvers import reverse
	from django.http import HttpResponse
	from django.shortcuts import get_object_or_404
	from django.shortcuts import render

	output = HttpResponse()

	# Je vérifie l'existence d'un objet TFinancement.
	o_fact = get_object_or_404(TFacture, pk = _f)

	# Je vérifie le droit de lecture.
	ger_droits(request.user, o_fact.id_doss)

	if request.method == 'GET' :

		# Je prépare le volet de consultation de la facture.
		t_attrs_fact = {
			'id_doss' : { 'label' : 'Numéro du dossier', 'value' : o_fact.id_doss },
			'id_prest' : { 'label' : 'Prestation', 'value' : o_fact.id_prest },
			'num_fact' : { 'label' : 'Numéro de facture', 'value' : o_fact.num_fact },
			'dt_mand_moa_fact' : {
				'label' : 'Date de mandatement par le maître d\'ouvrage',
				'value' : dt_fr(o_fact.dt_mand_moa_fact) or ''
			},
			'mont_ht_fact' : {
				'label' : 'Montant HT de la facture (en €)', 'value' : obt_mont(o_fact.mont_ht_fact) or ''
			},
			'mont_ttc_fact' : {
				'label' : 'Montant TTC de la facture (en €)', 'value' : obt_mont(o_fact.mont_ttc_fact) or ''
			},
			'dt_rec_fact' : { 'label' : 'Date de réception de la facture', 'value' : dt_fr(o_fact.dt_rec_fact) or '' },
			'num_mandat_fact' : { 'label' : 'Numéro de mandat', 'value' : o_fact.num_mandat_fact },
			'num_bord_fact' : { 'label' : 'Numéro de bordereau', 'value' : o_fact.num_bord_fact },
			'suivi_fact' : { 'label' : 'Suivi de la facturation', 'value' : o_fact.suivi_fact },
			'chem_pj_fact' : {
				'label' : 'Consulter le fichier scanné de la facture', 'value' : o_fact.chem_pj_fact, 'pdf' : True
			},
			'comm_fact' : { 'label' : 'Commentaire', 'value' : o_fact.comm_fact or '' },
		}

		# Je déclare un tableau de fenêtres modales.
		t_fm = [
			init_fm('suppr_fact', 'Supprimer une facture')
		]

		# J'affiche le template.
		output = render(
			request,
			'./gestion_dossiers/cons_fact.html',
			{
				'f' : o_fact,
				'forbidden' : ger_droits(request.user, o_fact.id_doss, False, False),
				't_attrs_fact' : init_pg_cons(t_attrs_fact),
				't_fm' : t_fm,
				'title' : 'Consulter une facture'
			}
		)

	else :
		if 'action' in request.GET :

			# Je traite le cas où je dois supprimer une facture.
			if request.GET['action'] == 'supprimer-facture' :
				output = HttpResponse(suppr(reverse('suppr_fact', args = [o_fact.pk])))

	return output

'''
Cette vue permet de traiter le formulaire d'ajout d'une demande de versement.
request : Objet requête
'''
@verif_acc
def ajout_ddv(request) :

	# Imports
	from app.forms.gestion_dossiers import GererDemandeVersement
	from app.functions import ger_droits
	from app.functions import init_f
	from app.functions import rempl_fich_log
	from app.models import TDossier
	from app.models import TFacture
	from app.models import TFacturesDemandeVersement
	from app.models import TFinancement
	from django.core.urlresolvers import reverse
	from django.http import HttpResponse
	import json

	output = HttpResponse()

	if request.method == 'POST' :
		if 'action' in request.GET :

			# Je stocke la valeur du paramètre "GET" dont la clé est "action".
			get_action = request.GET['action']

			# Je traite le cas où je dois filtrer les factures.
			if get_action == 'filtrer-factures' :

				# Tentative d'obtention d'une instance TFinancement
				try :
					o_fin = TFinancement.objects.get(pk = request.POST.get('GererDemandeVersement-zl_fin'))
				except :
					o_fin = None

				f_ajout_ddv = init_f(GererDemandeVersement(
					prefix = 'GererDemandeVersement',
					k_doss = TDossier.objects.get(num_doss = request.POST.get('GererDemandeVersement-za_num_doss')),
					k_fin = o_fin
				))
				output = HttpResponse(f_ajout_ddv['cbsm_fact'])

		else :

			# Je vérifie le droit d'écriture.
			try :
				o_doss_droit = TDossier.objects.get(num_doss = request.POST.get('GererDemandeVersement-za_num_doss'))
			except :
				o_doss_droit = None
			if o_doss_droit :
				ger_droits(request.user, o_doss_droit, False)

			# Tentative d'obtention d'une instance TFinancement
			try :
				o_fin = TFinancement.objects.get(pk = request.POST.get('GererDemandeVersement-zl_fin'))
			except :
				o_fin = None

			# Je soumets le formulaire.
			f_ajout_ddv = GererDemandeVersement(
				request.POST,
				request.FILES,
				prefix = 'GererDemandeVersement',
				k_fin = o_fin
			)

			# Je rajoute des choix valides pour les listes déroulantes des financements et des factures (prévention d'
			# erreurs).
			post_doss = request.POST.get('GererDemandeVersement-za_num_doss')
			post_fin = request.POST.get('GererDemandeVersement-zl_fin')
			if post_doss and post_fin :
				f_ajout_ddv.fields['zl_fin'].choices = [(f.pk, f.id_org_fin) for f in TFinancement.objects.filter(
					id_doss__num_doss = post_doss, pk = post_fin
				)]
			if post_fin :
				try :
					o_doss = TFinancement.objects.get(pk = post_fin).id_doss
				except :
					o_doss = None
				f_ajout_ddv.fields['cbsm_fact'].choices = [(f.pk, f) for f in TFacture.objects.filter(
					id_doss = o_doss
				)]

			if f_ajout_ddv.is_valid() :

				# Je créé la nouvelle instance TDemandeVersement.
				o_nvelle_ddv = f_ajout_ddv.save(commit = False)
				o_nvelle_ddv.save()

				# Je fais le lien avec la table TFacturesDemandeVersement.
				v_fact = request.POST.getlist('GererDemandeVersement-cbsm_fact')
				for i in range(0, len(v_fact)) :
					TFacturesDemandeVersement.objects.create(
						id_ddv = o_nvelle_ddv,
						id_fact = TFacture.objects.get(pk = v_fact[i])
					)

				# J'affiche le message de succès.
				output = HttpResponse(
					json.dumps({ 'success' : {
						'message' : '''
						La demande de versement a été ajoutée avec succès au partenaire financier « {0} » du dossier {1}.
						'''.format(o_nvelle_ddv.id_fin.id_org_fin, o_nvelle_ddv.id_fin.id_doss),
						'redirect' : reverse('cons_doss', args = [o_nvelle_ddv.id_fin.id_doss.pk])
					}}),
					content_type = 'application/json'
				)

				# Je renseigne l'onglet actif après rechargement de la page.
				request.session['tab_doss'] = '#ong_ddv'

				# Je complète le fichier log.
				rempl_fich_log([
					request.user.pk,
					request.user,
					o_nvelle_ddv.id_fin.id_doss.pk,
					o_nvelle_ddv.id_fin.id_doss,
					'C',
					'Demande de versement',
					o_nvelle_ddv.pk
				])

			else :

				# J'affiche les erreurs.
				t_err = {}
				for k, v in f_ajout_ddv.errors.items() :
					t_err['GererDemandeVersement-{0}'.format(k)] = v
				output = HttpResponse(json.dumps(t_err), content_type = 'application/json')

	return output

'''
Cette vue permet d'afficher la page de modification d'une demande de versement ou de traiter le formulaire de mise à
jour d'une demande de versement.
request : Objet requête
_d : Identifiant d'une demande de versement
'''
@verif_acc
@csrf_exempt
def modif_ddv(request, _d) :

	# Imports
	from app.forms.gestion_dossiers import GererDemandeVersement
	from app.functions import ger_droits
	from app.functions import init_f
	from app.functions import init_fm
	from app.functions import rempl_fich_log
	from app.models import TDemandeVersement
	from app.models import TFacture
	from app.models import TFacturesDemandeVersement
	from app.models import TFinancement
	from django.core.urlresolvers import reverse
	from django.http import HttpResponse
	from django.shortcuts import get_object_or_404
	from django.shortcuts import render
	import json

	output = HttpResponse()

	# Je vérifie l'existence d'un objet TDemandeVersement.
	o_ddv = get_object_or_404(TDemandeVersement, pk = _d)

	# Je vérifie le droit d'écriture.
	ger_droits(request.user, o_ddv.id_fin.id_doss, False)

	if request.method == 'GET' :

		# J'instancie un objet "formulaire".
		f_modif_ddv = GererDemandeVersement(
			instance = o_ddv,
			prefix = 'GererDemandeVersement',
			k_fin = o_ddv.id_fin,
			k_init = True
		)

		# Je déclare un tableau de fenêtres modales.
		t_fm = [
			init_fm('ger_ddv', 'Modifier une demande de versement')
		]

		# J'affiche le template.
		output = render(
			request,
			'./gestion_dossiers/modif_ddv.html',
			{
				'd' : o_ddv,
				'f_modif_ddv' : init_f(f_modif_ddv),
				't_fm' : t_fm,
				'title' : 'Modifier une demande de versement'
			}
		)

	else :
		if 'action' in request.GET :

			# Je stocke la valeur du paramètre "GET" dont la clé est "action".
			get_action = request.GET['action']

			# Je traite le cas où je dois filtrer les factures.
			if get_action == 'filtrer-factures' :

				# Tentative d'obtention d'une instance TFinancement
				try :
					o_fin = TFinancement.objects.get(pk = request.POST.get('GererDemandeVersement-zl_fin'))
				except :
					o_fin = None

				f_modif_ddv = init_f(
					GererDemandeVersement(
						instance = o_ddv,
						prefix = 'GererDemandeVersement',
						k_fin = o_fin
					)
				)
				output = HttpResponse(f_modif_ddv['cbsm_fact'])

		else :

			# Tentative d'obtention d'une instance TFinancement
			try :
				o_fin = TFinancement.objects.get(pk = request.POST.get('GererDemandeVersement-zl_fin'))
			except :
				o_fin = None

			# Je soumets le formulaire.
			f_modif_ddv = GererDemandeVersement(
				request.POST,
				request.FILES,
				instance = o_ddv,
				prefix = 'GererDemandeVersement',
				k_fin = o_fin
			)

			if f_modif_ddv.is_valid() :

				# Je modifie l'instance TDemandeVersement.
				o_ddv_modif = f_modif_ddv.save(commit = False)
				o_ddv_modif.save()

				# Je supprime tous les enregistrements liés à la demande de versement afin de les recréer.
				TFacturesDemandeVersement.objects.filter(id_ddv = o_ddv_modif).delete()

				# Je fais le lien avec la table TFacturesDemandeVersement.
				v_fact = request.POST.getlist('GererDemandeVersement-cbsm_fact')
				for i in range(0, len(v_fact)) :
					TFacturesDemandeVersement.objects.create(
						id_ddv = o_ddv_modif,
						id_fact = TFacture.objects.get(pk = v_fact[i])
					)

				# J'affiche le message de succès.
				output = HttpResponse(
					json.dumps({ 'success' : {
						'message' : 'La demande de versement a été mise à jour avec succès.',
						'redirect' : reverse('cons_ddv', args = [o_ddv_modif.pk])
					}}),
					content_type = 'application/json'
				)

				# Je complète le fichier log.
				rempl_fich_log([
					request.user.pk,
					request.user,
					o_ddv_modif.id_fin.id_doss.pk,
					o_ddv_modif.id_fin.id_doss,
					'U',
					'Demande de versement',
					o_ddv_modif.pk
				])

			else :

				# J'affiche les erreurs.
				t_err = {}
				for k, v in f_modif_ddv.errors.items() :
					t_err['GererDemandeVersement-{0}'.format(k)] = v
				output = HttpResponse(json.dumps(t_err), content_type = 'application/json')

	return output

'''
Cette vue permet de traiter le formulaire de suppression d'une demande de versement.
request : Objet requête
_d : Identifiant d'une demande de versement
'''
@verif_acc
@csrf_exempt
def suppr_ddv(request, _d) :

	# Imports
	from app.functions import ger_droits
	from app.functions import rempl_fich_log
	from app.models import TDemandeVersement
	from app.models import TFacturesDemandeVersement
	from django.core.urlresolvers import reverse
	from django.http import HttpResponse
	from django.shortcuts import get_object_or_404
	import json

	output = HttpResponse()

	# Je vérifie l'existence d'un objet TDemandeVersement.
	o_ddv = get_object_or_404(TDemandeVersement, pk = _d)

	if request.method == 'POST' :

		# Je vérifie le droit d'écriture.
		ger_droits(request.user, o_ddv.id_fin.id_doss, False)

		# Je supprime les enregistrements de la table TFacturesDemandeVersement relatifs à la demande de versement.
		TFacturesDemandeVersement.objects.filter(id_ddv = o_ddv).delete()

		# Je pointe vers l'objet TDemandeVersement à supprimer.
		o_ddv_suppr = o_ddv

		# Je récupère l'identifiant de la demande de versement afin de le renseigner dans le fichier log.
		v_id_ddv = o_ddv.pk

		# Je supprime l'objet TDemandeVersement.
		o_ddv.delete()

		# J'affiche le message de succès.
		output = HttpResponse(
			json.dumps({ 'success' : {
				'message' : 'La demande de versement a été supprimée avec succès.',
				'redirect' : reverse('cons_doss', args = [o_ddv_suppr.id_fin.id_doss.pk])
			}}),
			content_type = 'application/json'
		)

		# Je renseigne l'onglet actif après rechargement de la page.
		request.session['tab_doss'] = '#ong_ddv'

		# Je complète le fichier log.
		rempl_fich_log([
			request.user.pk,
			request.user,
			o_ddv_suppr.id_fin.id_doss.pk,
			o_ddv_suppr.id_fin.id_doss,
			'D',
			'Demande de versement',
			v_id_ddv
		])

	return output

'''
Cette vue permet d'afficher la page de consultation d'une demande de versement.
request : Objet requête
_d : Identifiant d'une demande de versement
'''
@verif_acc
@csrf_exempt
def cons_ddv(request, _d) :

	# Imports
	from app.functions import dt_fr
	from app.functions import ger_droits
	from app.functions import init_fm
	from app.functions import init_pg_cons
	from app.functions import obt_mont
	from app.functions import suppr
	from app.models import TDemandeVersement
	from app.models import TFacturesDemandeVersement
	from app.models import VDemandeVersement
	from django.core.urlresolvers import reverse
	from django.http import HttpResponse
	from django.shortcuts import get_object_or_404
	from django.shortcuts import render

	output = HttpResponse()

	# Je vérifie l'existence d'un objet TDemandeVersement.
	o_ddv = get_object_or_404(TDemandeVersement, pk = _d)

	if request.method == 'GET' :

		# Je vérifie le droit de lecture.
		ger_droits(request.user, o_ddv.id_fin.id_doss)

		o_suivi_ddv = VDemandeVersement.objects.get(pk = o_ddv.pk)

		# Je prépare le volet de consultation de la demande de versement.
		t_attrs_ddv = {
			'id_doss' : { 'label' : 'Numéro du dossier', 'value' : o_suivi_ddv.id_doss },
			'id_org_fin' : { 'label' : 'Partenaire financier', 'value' : o_suivi_ddv.id_org_fin },
			'id_type_vers' : { 'label' : 'Type de demande de versement', 'value' : o_ddv.id_type_vers },
			'int_ddv' : { 'label' : 'Intitulé de la demande de versement', 'value' : o_ddv.int_ddv },
			'mont_ht_ddv' : {
				'label' : 'Montant HT de la demande de versement (en €)', 'value' : obt_mont(o_ddv.mont_ht_ddv) or ''
			},
			'mont_ttc_ddv' : {
				'label' : 'Montant TTC de la demande de versement (en €)', 'value' : obt_mont(o_ddv.mont_ttc_ddv) or ''
			},
			'dt_ddv' : { 'label' : 'Date de la demande de versement', 'value' : dt_fr(o_ddv.dt_ddv) },
			'dt_vers_ddv' : { 'label' : 'Date de versement', 'value' : dt_fr(o_ddv.dt_vers_ddv) or '' },
			'num_bord_ddv' : { 'label' : 'Numéro de bordereau', 'value' : o_ddv.num_bord_ddv },
			'num_titre_rec_ddv' : { 'label' : 'Numéro de titre de recette', 'value' : o_ddv.num_titre_rec_ddv },
			'mont_ht_verse_ddv' : {
				'label' : 'Montant HT versé (en €)', 'value' : obt_mont(o_ddv.mont_ht_verse_ddv) or ''
			},
			'mont_ttc_verse_ddv' : {
				'label' : 'Montant TTC versé (en €)', 'value' : obt_mont(o_ddv.mont_ttc_verse_ddv) or ''
			},
			'map_ht_ddv' : { 'label' : 'Manque à percevoir HT (en €)', 'value' : obt_mont(o_suivi_ddv.map_ht_ddv) or '' },
			'map_ttc_ddv' : { 'label' : 'Manque à percevoir TTC (en €)', 'value' : obt_mont(o_suivi_ddv.map_ttc_ddv) or '' },
			'chem_pj_ddv' : {
				'label' : 'Consulter le courrier scanné de la demande de versement',
				'value' : o_ddv.chem_pj_ddv,
				'pdf' : True
			},
			'comm_ddv' : { 'label' : 'Commentaire', 'value' : o_ddv.comm_ddv or '' }
		}

		# J'initialise le tableau des factures reliées à la demande de versement.
		t_fact_ddv = []
		for fd in TFacturesDemandeVersement.objects.filter(id_ddv = o_ddv) :

			mont_fact = fd.id_fact.mont_ht_fact
			if fd.id_fact.id_doss.est_ttc_doss == True :
				mont_fact = fd.id_fact.mont_ttc_fact

			t_fact_ddv.append({
				'id_prest' : fd.id_fact.id_prest,
				'num_fact' : fd.id_fact.num_fact,
				'dt_mand_moa_fact' : dt_fr(fd.id_fact.dt_mand_moa_fact) or '-',
				'mont_fact' : obt_mont(mont_fact),
				'id_fact__pk' : fd.id_fact.pk
			})

		# Je définis si le montant du dossier est en HT ou en TTC.
		ht_ou_ttc = 'HT'
		if o_ddv.id_fin.id_doss.est_ttc_doss == True :
			ht_ou_ttc = 'TTC'

		# Je déclare un tableau de fenêtres modales.
		t_fm = [
			init_fm('suppr_ddv', 'Supprimer une demande de versement')
		]

		# J'affiche le template.
		output = render(
			request,
			'./gestion_dossiers/cons_ddv.html',
			{
				'd' : o_ddv,
				'forbidden' : ger_droits(request.user, o_ddv.id_fin.id_doss, False, False),
				'ht_ou_ttc' : ht_ou_ttc,
				't_attrs_ddv' : init_pg_cons(t_attrs_ddv),
				't_fact_ddv' : t_fact_ddv,
				't_fm' : t_fm,
				'title' : 'Consulter une demande de versement'
			}
		)

	else :
		if 'action' in request.GET :

			# Je traite le cas où je dois supprimer une demande de versement.
			if request.GET['action'] == 'supprimer-demande-versement' :
				output = HttpResponse(suppr(reverse('suppr_ddv', args = [o_ddv.pk])))

	return output

'''
Cette vue permet de traiter le formulaire d'ajout d'un arrêté.
request : Objet requête
'''
@verif_acc
def ajout_arr(request) :

	# Imports
	from app.forms.gestion_dossiers import GererArrete
	from app.functions import ger_droits
	from app.functions import rempl_fich_log
	from app.models import TDossier
	from django.core.urlresolvers import reverse
	from django.http import HttpResponse
	import json

	output = HttpResponse()

	if request.method == 'POST' :

		# Je vérifie le droit d'écriture.
		try :
			o_doss_droit = TDossier.objects.get(num_doss = request.POST.get('GererArrete-za_num_doss'))
		except :
			o_doss_droit = None
		if o_doss_droit :
			ger_droits(request.user, o_doss_droit, False)

		# Je soumets le formulaire.
		f_ajout_arr = GererArrete(request.POST, request.FILES, prefix = 'GererArrete')

		if f_ajout_arr.is_valid() :

			# Je créé la nouvelle instance TArretesDossier.
			o_nv_arr = f_ajout_arr.save(commit = False)
			o_nv_arr.save()

			# J'affiche le message de succès.
			output = HttpResponse(
				json.dumps({ 'success' : {
					'message' : 'Le type de déclaration « {0} » a été ajouté avec succès au dossier {1}.'.format(
						o_nv_arr.id_type_decl, o_nv_arr.id_doss
					),
					'redirect' : reverse('cons_doss', args = [o_nv_arr.id_doss.pk])
				}}),
				content_type = 'application/json'
			)

			# Je renseigne l'onglet actif après rechargement de la page.
			request.session['tab_doss'] = '#ong_arr'

			# Je complète le fichier log.
			rempl_fich_log([
				request.user.pk, request.user, o_nv_arr.id_doss.pk, o_nv_arr.id_doss, 'C', 'Arrêté', o_nv_arr.pk
			])

		else :

			# J'affiche les erreurs.
			t_err = {}
			for k, v in f_ajout_arr.errors.items() :
				t_err['GererArrete-{0}'.format(k)] = v
			output = HttpResponse(json.dumps(t_err), content_type = 'application/json')

	return output

'''
Cette vue permet d'afficher la page de modification d'un arrêté ou de traiter le formulaire de mise à jour d'un arrêté.
request : Objet requête
_a : Identifiant d'un arrêté
'''
@verif_acc
def modif_arr(request, _a) :

	# Imports
	from app.forms.gestion_dossiers import GererArrete
	from app.functions import ger_droits
	from app.functions import init_f
	from app.functions import init_fm
	from app.functions import rempl_fich_log
	from app.models import TArretesDossier
	from django.core.urlresolvers import reverse
	from django.http import HttpResponse
	from django.shortcuts import get_object_or_404
	from django.shortcuts import render
	import json

	output = HttpResponse()

	# Je vérifie l'existence d'un objet TArretesDossier
	o_arr = get_object_or_404(TArretesDossier, pk = _a)

	# Je vérifie le droit d'écriture.
	ger_droits(request.user, o_arr.id_doss, False)

	if request.method == 'GET' :

		# J'instancie un objet "formulaire".
		f_modif_arr = GererArrete(instance = o_arr, prefix = 'GererArrete')

		# Je déclare un tableau de fenêtres modales.
		t_fm = [
			init_fm('modif_arr', 'Modifier un arrêté')
		]

		# J'affiche le template.
		output = render(
			request,
			'./gestion_dossiers/modif_arr.html',
			{ 'a' : o_arr, 'f_modif_arr' : init_f(f_modif_arr), 't_fm' : t_fm, 'title' : 'Modifier un arrêté' }
		)

	else :

		# Je soumets le formulaire.
		f_modif_arr = GererArrete(request.POST, request.FILES, instance = o_arr, prefix = 'GererArrete')

		if f_modif_arr.is_valid() :

			# Je modifie l'instance TArretesDossier.
			o_arr_modif = f_modif_arr.save(commit = False)
			o_arr_modif.save()

			# J'affiche le message de succès.
			output = HttpResponse(
				json.dumps({ 'success' : {
					'message' : 'Le type de déclaration « {0} » du dossier {1} a été mis à jour avec succès.'.format(
						o_arr_modif.id_type_decl, o_arr_modif.id_doss
					),
					'redirect' : reverse('cons_arr', args = [o_arr_modif.pk])
				}}),
				content_type = 'application/json'
			)

			# Je complète le fichier log.
			rempl_fich_log([
				request.user.pk,
				request.user,
				o_arr_modif.id_doss.pk,
				o_arr_modif.id_doss,
				'U',
				'Arrêté',
				o_arr_modif.pk
			])

		else :

			# J'affiche les erreurs.
			t_err = {}
			for k, v in f_modif_arr.errors.items() :
				t_err['GererArrete-{0}'.format(k)] = v
			output = HttpResponse(json.dumps(t_err), content_type = 'application/json')

	return output

'''
Cette vue permet de traiter le formulaire de suppression d'un arrêté.
request : Objet requête
_a : Identifiant d'un arrêté
'''
@verif_acc
@csrf_exempt
def suppr_arr(request, _a) :

	# Imports
	from app.functions import ger_droits
	from app.functions import rempl_fich_log
	from app.models import TArretesDossier
	from django.core.urlresolvers import reverse
	from django.http import HttpResponse
	from django.shortcuts import get_object_or_404
	import json

	output = HttpResponse()

	# Je vérifie l'existence d'un objet TArretesDossier.
	o_arr = get_object_or_404(TArretesDossier, pk = _a)

	if request.method == 'POST' :

		# Je vérifie le droit d'écriture.
		ger_droits(request.user, o_arr.id_doss, False)

		# Je pointe vers les objets TDossier et TTypeDeclaration liés à l'objet TArretesDossier.
		o_doss = o_arr.id_doss
		o_type_decl = o_arr.id_type_decl

		# Je récupère l'identifiant de l'arrêté afin de le renseigner dans le fichier log.
		v_id_arr_doss = o_arr.pk

		# Je supprime l'objet TArretesDossier.
		o_arr.delete()

		# J'affiche le message de succès.
		output = HttpResponse(
			json.dumps({ 'success' : {
				'message' : 'Le type de déclaration « {0} » a été supprimé avec succès du dossier {1}.'.format(
					o_type_decl, o_doss
				),
				'redirect' : reverse('cons_doss', args = [o_doss.pk])
			}}),
			content_type = 'application/json'
		)

		# Je renseigne l'onglet actif après rechargement de la page.
		request.session['tab_doss'] = '#ong_arr'

		# Je complète le fichier log.
		rempl_fich_log([request.user.pk, request.user, o_doss.pk, o_doss, 'D', 'Arrêté', v_id_arr_doss])

	return output

'''
Cette vue permet d'afficher la page de consultation d'un arrêté.
request : Objet requête
_a : Identifiant d'un arrêté
'''
@verif_acc
@csrf_exempt
def cons_arr(request, _a) :

	# Imports
	from app.functions import dt_fr
	from app.functions import ger_droits
	from app.functions import init_fm
	from app.functions import init_pg_cons
	from app.functions import suppr
	from app.models import TArretesDossier
	from django.core.urlresolvers import reverse
	from django.http import HttpResponse
	from django.shortcuts import get_object_or_404
	from django.shortcuts import render

	output = HttpResponse()

	# Je vérifie l'existence d'un objet TArretesDossier.
	o_arr = get_object_or_404(TArretesDossier, pk = _a)

	if request.method == 'GET' :

		# Je vérifie le droit de lecture.
		ger_droits(request.user, o_arr.id_doss)

		# Je prépare le volet de consultation de l'arrêté.
		t_attrs_arr = {
			'id_doss' : { 'label' : 'Numéro du dossier', 'value' : o_arr.id_doss },
			'id_type_decl' : { 'label' : 'Type de déclaration', 'value' : o_arr.id_type_decl },
			'id_type_av_arr' : { 'label' : 'Avancement', 'value' : o_arr.id_type_av_arr },
			'num_arr' : { 'label' : 'Numéro de l\'arrêté', 'value' : o_arr.num_arr },
			'dt_sign_arr' : { 'label' : 'Date de signature de l\'arrêté', 'value' : dt_fr(o_arr.dt_sign_arr) or '' },
			'dt_lim_encl_trav_arr' : {
				'label' : 'Date limite d\'enclenchement des travaux', 'value' : dt_fr(o_arr.dt_lim_encl_trav_arr) or ''
			},
			'chem_pj_arr' : {
				'label' : 'Consulter le fichier scanné de l\'arrêté', 'value' : o_arr.chem_pj_arr, 'pdf' : True
			},
			'comm_arr' : { 'label' : 'Commentaire', 'value' : o_arr.comm_arr or '' },
		}

		# Je déclare un tableau de fenêtres modales.
		t_fm = [
			init_fm('suppr_arr', 'Supprimer un arrêté')
		]

		# J'affiche le template.
		output = render(
			request,
			'./gestion_dossiers/cons_arr.html',
			{
				'a' : o_arr,
				'forbidden' : ger_droits(request.user, o_arr.id_doss, False, False),
				't_attrs_arr' : init_pg_cons(t_attrs_arr),
				't_fm' : t_fm,
				'title' : 'Consulter un arrêté'
			}
		)

	else :
		if 'action' in request.GET :

			# Je traite le cas où je dois supprimer un arrêté.
			if request.GET['action'] == 'supprimer-arrete' :
				if request.GET['arrete'] :
					output = HttpResponse(suppr(reverse('suppr_arr', args = [request.GET['arrete']])))

	return output

'''
Cette vue permet de traiter le formulaire d'ajout d'une photo.
request : Objet requête
'''
@verif_acc
def ajout_ph(request) :

	# Imports
	from app.forms.gestion_dossiers import GererPhoto
	from app.functions import ger_droits
	from app.functions import rempl_fich_log
	from app.models import TDossier
	from django.core.urlresolvers import reverse
	from django.http import HttpResponse
	import json

	output = HttpResponse()

	if request.method == 'POST' :

		# Je vérifie le droit d'écriture.
		try :
			o_doss_droit = TDossier.objects.get(num_doss = request.POST.get('GererPhoto-za_num_doss'))
		except :
			o_doss_droit = None
		if o_doss_droit :
			ger_droits(request.user, o_doss_droit, False)

		# Je soumets le formulaire.
		f_ajout_ph = GererPhoto(request.POST, request.FILES, prefix = 'GererPhoto')

		if f_ajout_ph.is_valid() :

			# Je créé la nouvelle instance TPhoto.
			o_nvelle_ph = f_ajout_ph.save(commit = False)
			o_nvelle_ph.save()

			# J'affiche le message de succès.
			output = HttpResponse(
				json.dumps({ 'success' : {
					'message' : 'La photo a été ajoutée avec succès au dossier {0}.'.format(o_nvelle_ph.id_doss),
					'redirect' : reverse('cons_doss', args = [o_nvelle_ph.id_doss.pk])
				}}),
				content_type = 'application/json'
			)

			# Je renseigne l'onglet actif après rechargement de la page.
			request.session['tab_doss'] = '#ong_ph'

			# Je complète le fichier log.
			rempl_fich_log([
				request.user.pk,
				request.user,
				o_nvelle_ph.id_doss.pk,
				o_nvelle_ph.id_doss,
				'C',
				'Photo',
				o_nvelle_ph.pk
			])

		else :

			# J'affiche les erreurs.
			t_err = {}
			for k, v in f_ajout_ph.errors.items() :
				t_err['GererPhoto-{0}'.format(k)] = v
			output = HttpResponse(json.dumps(t_err), content_type = 'application/json')

	return output

'''
Cette vue permet d'afficher la page de modification d'une photo ou de traiter le formulaire de mise à jour
d'une photo.
request : Objet requête
_p : Identifiant d'une photo
'''
@verif_acc
def modif_ph(request, _p) :

	# Imports
	from app.forms.gestion_dossiers import GererPhoto
	from app.functions import ger_droits
	from app.functions import init_f
	from app.functions import init_fm
	from app.functions import rempl_fich_log
	from django.core.urlresolvers import reverse
	from app.models import TPhoto
	from django.http import HttpResponse
	from django.shortcuts import get_object_or_404
	from django.shortcuts import render
	import json

	output = HttpResponse()

	# Je vérifie l'existence d'un objet TPhoto.
	o_ph = get_object_or_404(TPhoto, pk = _p)

	# Je vérifie le droit d'écriture.
	ger_droits(request.user, o_ph.id_doss, False)

	if request.method == 'GET' :

		# J'instancie un objet "formulaire".
		f_modif_ph = GererPhoto(instance = o_ph, prefix = 'GererPhoto')

		# Je déclare un tableau de fenêtres modales.
		t_fm = [
			init_fm('modif_ph', 'Modifier une photo')
		]

		# J'affiche le template.
		output = render(
			request,
			'./gestion_dossiers/modif_ph.html',
			{ 'f_modif_ph' : init_f(f_modif_ph), 'p' : o_ph, 't_fm' : t_fm, 'title' : 'Modifier une photo' }
		)

	else :

		# Je soumets le formulaire.
		f_modif_ph = GererPhoto(request.POST, request.FILES, instance = o_ph, prefix = 'GererPhoto')

		if f_modif_ph.is_valid() :

			# Je modifie l'instance TPhoto.
			o_ph_modif = f_modif_ph.save(commit = False)
			o_ph_modif.save()

			# J'affiche le message de succès.
			output = HttpResponse(
				json.dumps({ 'success' : {
					'message' : 'La photo a été mise à jour avec succès.',
					'redirect' : reverse('cons_doss', args = [o_ph.id_doss.pk])
				}}),
				content_type = 'application/json'
			)

			# Je renseigne l'onglet actif après rechargement de la page.
			request.session['tab_doss'] = '#ong_ph'

			# Je complète le fichier log.
			rempl_fich_log([
				request.user.pk, request.user, o_ph_modif.id_doss.pk, o_ph_modif.id_doss, 'U', 'Photo', o_ph_modif.pk
			])

		else :

			# J'affiche les erreurs.
			t_err = {}
			for k, v in f_modif_ph.errors.items() :
				t_err['GererPhoto-{0}'.format(k)] = v
			output = HttpResponse(json.dumps(t_err), content_type = 'application/json')

	return output

'''
Cette vue permet de traiter le formulaire de suppression d'une photo.
request : Objet requête
_p : Identifiant d'une photo
'''
@verif_acc
@csrf_exempt
def suppr_ph(request, _p) :

	# Imports
	from app.functions import ger_droits
	from app.functions import rempl_fich_log
	from app.models import TPhoto
	from django.core.urlresolvers import reverse
	from django.http import HttpResponse
	from django.shortcuts import get_object_or_404
	import json

	output = HttpResponse()

	# Je vérifie l'existence d'un objet TPhoto.
	o_ph = get_object_or_404(TPhoto, pk = _p)

	if request.method == 'POST' :

		# Je vérifie le droit d'écriture.
		ger_droits(request.user, o_ph.id_doss, False)

		# Je pointe vers l'objet TDossier lié à l'objet TPhoto.
		o_doss = o_ph.id_doss

		# Je récupère l'identifiant de la photo afin de le renseigner dans le fichier log.
		v_id_ph = o_ph.pk

		# Je supprime l'objet TPhoto.
		o_ph.delete()

		# J'affiche le message de succès.
		output = HttpResponse(
			json.dumps({ 'success' : {
				'message' : 'La photo a été supprimée avec succès du dossier {0}.'.format(o_doss),
				'redirect' : reverse('cons_doss', args = [o_doss.pk])
			}}),
			content_type = 'application/json'
		)

		# Je renseigne l'onglet actif après rechargement de la page.
		request.session['tab_doss'] = '#ong_ph'

		# Je complète le fichier log.
		rempl_fich_log([request.user.pk, request.user, o_doss.pk, o_doss, 'D', 'Photo', v_id_ph])

	return output

'''
Cette vue permet de traiter le formulaire d'ajout d'un prestataire.
request : Objet requête
'''
@verif_acc
def ajout_org_prest(request) :

	# Imports
	from app.forms.gestion_dossiers import AjouterPrestataire
	from app.functions import rempl_fich_log
	from django.http import HttpResponse
	import json

	output = HttpResponse()

	if request.method == 'POST' :

		# Je soumets le formulaire.
		f_ajout_org_prest = AjouterPrestataire(request.POST, prefix = 'AjouterPrestataire')

		if f_ajout_org_prest.is_valid() :

			# Je créé la nouvelle instance TPrestataire.
			o_nv_org_prest = f_ajout_org_prest.save(commit = False)
			o_nv_org_prest.save()

			# J'affiche le message de succès.
			output = HttpResponse(
				json.dumps({ 'success' : {
					'message' : 'Le prestataire {0} a été créé avec succès.'.format(o_nv_org_prest),
					'display' : ['#id_GererPrestation-zs_siret_org_prest', o_nv_org_prest.siret_org_prest]
				}}),
				content_type = 'application/json'
			)

			# Je complète le fichier log.
			rempl_fich_log([request.user.pk, request.user, None, None, 'C', 'Prestataire', o_nv_org_prest.pk])

		else :

			# J'affiche les erreurs.
			t_err = {}
			for k, v in f_ajout_org_prest.errors.items() :
				t_err['AjouterPrestataire-{0}'.format(k)] = v
			output = HttpResponse(json.dumps(t_err), content_type = 'application/json')

	return output

'''
Cette vue permet d'afficher la page relative à un dossier qui sera par la suite convertie en PDF.
request : Objet requête
_d : Identifiant d'un dossier
'''
@verif_acc
def impr_doss(request, _d) :

	# Imports
	from app.functions import dt_fr
	from app.functions import ger_droits
	from app.functions import init_pg_cons
	from app.functions import obt_mont
	from app.functions import obt_pourc
	from app.forms.gestion_dossiers import PrintDoss
	from app.models import TDossier
	from app.models import TFacture
	from app.models import VDemandeVersement
	from app.models import VFinancement
	from app.models import VSuiviDossier
	from app.models import VSuiviPrestationsDossier
	from django.http import HttpResponse
	from django.shortcuts import get_object_or_404
	from django.shortcuts import render

	output = HttpResponse()

	# Je vérifie l'existence d'un objet TDossier.
	o_doss = get_object_or_404(TDossier, pk = _d)

	# Je vérifie le droit de lecture.
	ger_droits(request.user, o_doss)

	# Je pointe vers l'objet VSuiviDossier.
	o_suivi_doss = VSuiviDossier.objects.get(pk = o_doss.pk)

	# Je définis le mode de taxe du dossier.
	ht_ou_ttc = o_doss.est_ttc_doss and 'TTC' or 'HT'

	# Obtention des différents tableaux déployés dans chaque onglet
	doss_fam = o_doss.get_doss_fam()
	fdvs = o_doss.get_recap_fdvs()
	fins = o_doss.get_recap_fins()
	prss = o_doss.get_recap_prss()
	facs = o_doss.get_recap_facs()
	ddvs = o_doss.get_recap_ddvs()
	form = PrintDoss()
	context = {
		'd' : o_doss,
		'ht_ou_ttc' : ht_ou_ttc,
		'mont_ddv_sum_str' : obt_mont(ddvs['mnt']),
		'mont_doss' : obt_mont(o_doss.mont_doss),
		'mont_fact_sum' : facs['mnt'],
		'mont_ht_fact_sum' : facs['mnt_ht'],
		'mont_ttc_fact_sum' : facs['mnt_ttc'],
		'mont_fact_sum_str' : obt_mont(facs['mnt']),
		'mont_tot_prest_doss' : obt_mont(o_suivi_doss.mont_tot_prest_doss),
		'mont_rae' : obt_mont(o_suivi_doss.mont_rae),
		'mont_suppl_doss' : obt_mont(o_doss.mont_suppl_doss),
		'pourc_comm' : obt_pourc(o_suivi_doss.pourc_comm),
		'pourc_paye' : obt_pourc(o_suivi_doss.pourc_paye),
		't_attrs_doss' : init_pg_cons(o_doss.get_attrs(), True),
		't_ddv' : ddvs['tbl'],
		't_ddv_length' : len(ddvs['tbl']),
		't_doss_fam' : doss_fam,
		't_doss_fam_length' : len(doss_fam),
		't_fact' : facs['tbl'],
		't_fact_length' : len(facs['tbl']),
		't_fdvs' : fdvs,
		't_fin' : fins,
		't_prest' : prss['tbl'],
		't_prest_length' : len(prss['tbl']),
		't_prest_sum' : prss['mnts'],
		'features': {ff:True for ff in form.fields}
	}

	# Dans le cas d'un simple GET on retourne l'ensemble des sections
	if request.method == 'GET':

		output = render(
			request,
			'./gestion_dossiers/impr_doss.html',
			context=context
		)

	# Dans le cas d'un POST on va contextualiser le gabarit
	if request.method == 'POST':
		form = PrintDoss(request.POST)
		if form.is_valid():
			context.pop('features', None)
			context['features'] = form.cleaned_data

		output = render(
			request,
			'./gestion_dossiers/impr_doss.html',
			context=context
		)

	return output

'''
Cette vue permet le téléchargement d'une lettre type de demande de versement au format DOCX.
request : Objet requête
_d : Identifiant d'une demande de versement
'''
def edit_lt_ddv(request, _d) :

	# Imports
	from app.functions import dt_fr
	from app.functions import gen_cdc
	from app.functions import ger_droits
	from app.functions import obt_mont
	from app.models import TDemandeVersement
	from app.models import TFacturesDemandeVersement
	from docx import Document
	from docx.enum.text import WD_ALIGN_PARAGRAPH
	from docx.oxml import parse_xml
	from docx.oxml.ns import nsdecls
	from docx.shared import Mm
	from docx.shared import Pt
	from django.http import HttpResponse
	from django.shortcuts import get_object_or_404
	from styx.settings import MEDIA_ROOT
	from styx.settings import T_DONN_BDD_STR
	import time

	output = HttpResponse()

	# Je vérifie l'existence d'un objet TDemandeVersement.
	obj_ddv = get_object_or_404(TDemandeVersement, pk = _d)

	# Je vérifie le droit d'écriture.
	ger_droits(request.user, obj_ddv.id_fin.id_doss, False)

	if request.method == 'GET' :

		# Déclaration d'un nouveau document Word
		document = Document()

		# Définition de la police et de sa taille
		style = document.styles['Normal']
		font = style.font
		font.name = 'Calibri'
		font.size = Pt(10)

		# Obtention d'un objet TDossier
		obj_doss = obj_ddv.id_fin.id_doss

		# Obtention d'un objet TMoa
		obj_org_moa = obj_ddv.id_fin.id_doss.id_org_moa

		# Déclaration du tableau "en-tête"
		table = document.add_table(rows = 0, cols = 2)
		row_cells = table.add_row().cells

		# Insertion du logo du maître d'ouvrage
		paragraph = row_cells[0].paragraphs[0]
		if obj_org_moa.logo_org_moa :
			run = paragraph.add_run().add_picture('{0}/{1}'.format(MEDIA_ROOT, obj_org_moa.logo_org_moa), Mm(40))
		else :
			run = paragraph.add_run('[LOGO]')

		# Insertion de l'adresse complète du maître d'ouvrage
		paragraph = row_cells[1].paragraphs[0]
		run1 = paragraph.add_run(str(obj_org_moa))
		run1.bold = True
		run1.add_break()
		run2 = paragraph.add_run('{0} - {1} {2}'.format(
			obj_org_moa.adr_org or '[ADRESSE]',
			obj_org_moa.cp_org or '[CODE POSTAL]',
			obj_org_moa.num_comm.n_comm if obj_org_moa.num_comm else '[COMMUNE]'
		))
		run2.add_break()
		run3 = paragraph.add_run('Tél :')
		run3.underline = True
		run4 = paragraph.add_run(' {} - '.format(obj_org_moa.tel_org or '[TÉLÉPHONE]'))
		run5 = paragraph.add_run('Email :')
		run5.underline = True
		run6 = paragraph.add_run(' {}'.format(obj_org_moa.courr_org or '[EMAIL]'))

		# Mise en forme de la cellule.
		for r in [run1, run2, run3, run4, run5, run6] :
			r.font.size = Pt(8)
		paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

		# Nettoyage de certaines variables (précaution)
		del run1, run2, run3, run4, run5, run6

		# Obtention d'un objet TFinancement
		obj_fin = obj_ddv.id_fin

		# Insertion de certaines données liées au financement et au dossier
		paragraph = document.add_paragraph()
		run1 = paragraph.add_run().add_break()
		run2 = paragraph.add_run('Financeur :')
		run2.underline = True
		run3 = paragraph.add_run(' {}'.format(obj_fin.id_org_fin))
		run3.add_break()
		run4 = paragraph.add_run('N° de l\'arrêté ou convention :')
		run4.underline = True
		run5 = paragraph.add_run(' {}'.format(obj_fin.num_arr_fin or '[NUMÉRO DE L\'ARRÊTÉ OU CONVENTION]'))
		run5.add_break()
		run6 = paragraph.add_run('Taux de l\'aide :')
		run6.underline = True
		run7 = paragraph.add_run(
			' {} % (pourcentage éligible)'.format(obj_fin.pourc_elig_fin) if obj_fin.pourc_elig_fin else ' Aucun'
		)
		run8 = paragraph.add_run().add_break()
		run9 = paragraph.add_run().add_break()
		run10 = paragraph.add_run('Numéro du dossier SMMAR :')
		run10.underline = True
		run11 = paragraph.add_run(' {}'.format(obj_doss))
		run11.add_break()
		run12 = paragraph.add_run('Intitulé du dossier SMMAR :')
		run12.underline = True
		run13 = paragraph.add_run(' {}'.format(obj_doss.get_int_doss()))

		# Mise en forme du paragraphe
		paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

		# Nettoyage de certaines variables (précaution)
		del run1, run2, run3, run4, run5, run6, run7, run8, run9, run10, run11, run12, run13

		# Détermination du numéro de la demande de versment pour le couple dossier/financeur
		i = None
		qs_ddv = TDemandeVersement.objects.filter(
			id_fin__id_doss = obj_doss, id_fin__id_org_fin = obj_fin.id_org_fin
		).order_by('dt_ddv')
		for index, d in enumerate(qs_ddv) :
			if d == obj_ddv :
				i = index + 1

		# Insertion du titre de la demande de versement
		paragraph = document.add_paragraph()
		run = paragraph.add_run('Demande de versement de subvention n°{0} : {1}'.format(
			i, str(obj_ddv.id_type_vers).lower())
		)

		# Initialisation des déterminants démonstratifs
		tab_determ = {
			T_DONN_BDD_STR['TVERS_ACOMPT'] : 'cet {}'.format(T_DONN_BDD_STR['TVERS_ACOMPT'].lower()),
			T_DONN_BDD_STR['TVERS_AF'] : 'cette {}'.format(T_DONN_BDD_STR['TVERS_AF'].lower()),
			T_DONN_BDD_STR['TVERS_SOLDE'] : 'ce {}'.format(T_DONN_BDD_STR['TVERS_SOLDE'].lower())
		}

		# Mise en forme du paragraphe
		run.bold = True
		run.underline = True
		paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

		# Calcul du montant total des factures composant la demande de versement
		qs_fact_ddv = TFacturesDemandeVersement.objects.filter(id_ddv = obj_ddv).order_by('id_fact')
		mont_ht_fact_sum = 0
		mont_ttc_fact_sum = 0
		for fddv in qs_fact_ddv :
			mont_ht_fact_sum += fddv.id_fact.mont_ht_fact or 0
			mont_ttc_fact_sum += fddv.id_fact.mont_ttc_fact or 0

		# Insertion du montant total des factures si besoin
		if obj_ddv.id_type_vers.int_type_vers != T_DONN_BDD_STR['TVERS_AF'] :
			paragraph = document.add_paragraph('Veuillez trouver ci-joint la demande de versement n°{0} correspondant à un montant de factures de {1} € {2} pour {3}, tel que présenté dans le tableau récapitulatif ci-dessous.'.format(
				i,
				mont_ttc_fact_sum if obj_doss.est_ttc_doss == True else mont_ht_fact_sum,
				'TTC' if obj_doss.est_ttc_doss == True else 'HT',
				tab_determ[obj_ddv.id_type_vers.int_type_vers]
			))

		# Insertion du montant de l'aide
		paragraph = document.add_paragraph('L\'aide attendue pour {0} est de {1} € {2}.'.format(
			tab_determ[obj_ddv.id_type_vers.int_type_vers],
			obj_ddv.mont_ttc_ddv if obj_doss.est_ttc_doss == True else obj_ddv.mont_ht_ddv,
			'TTC' if obj_doss.est_ttc_doss == True else 'HT'
		))

		if obj_ddv.id_type_vers.int_type_vers != T_DONN_BDD_STR['TVERS_AF'] :

			# Déclaration du tableau des factures
			table = document.add_table(rows = 1, cols = 7, style = 'TableGrid')

			# Préparation de l'en-tête du tableau des factures
			header_cells = table.rows[0].cells
			paragraph1 = header_cells[0].paragraphs[0]
			run1 = paragraph1.add_run('Prestataire')
			paragraph2 = header_cells[1].paragraphs[0]
			run2 = paragraph2.add_run('Objet')
			paragraph3 = header_cells[2].paragraphs[0]
			run3 = paragraph3.add_run('Date de mandatement par le maître d\'ouvrage')
			paragraph4 = header_cells[3].paragraphs[0]
			run4 = paragraph4.add_run('N° de bordereau')
			paragraph5 = header_cells[4].paragraphs[0]
			run5 = paragraph5.add_run('N° de mandat')
			paragraph6 = header_cells[5].paragraphs[0]
			run6 = paragraph6.add_run('Montant HT (en €)')
			paragraph7 = header_cells[6].paragraphs[0]
			run7 = paragraph7.add_run('Montant TTC (en €)')

			# Mise en forme de l'en-tête du tableau des factures
			for r in [run1, run2, run3, run4, run5, run6, run7] :
				r.font.size = Pt(8)
				r.bold = True
			for p in [paragraph1, paragraph2, paragraph3, paragraph4, paragraph5, paragraph6, paragraph7] :
				p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

			# Nettoyage de certaines variables (précaution)
			del paragraph1, paragraph2, paragraph3, paragraph4, paragraph5, paragraph6, paragraph7
			del run1, run2, run3, run4, run5, run6, run7

			# Remplissage du tableau des factures
			for fddv in qs_fact_ddv :

				# Obtention d'une instance TFacture
				obj_fact = fddv.id_fact

				# Préparation d'une ligne du tableau des factures
				row_cells = table.add_row().cells
				paragraph1 = row_cells[0].paragraphs[0]
				run1 = paragraph1.add_run(str(obj_fact.id_prest.id_org_prest))
				paragraph2 = row_cells[1].paragraphs[0]
				run2 = paragraph2.add_run(obj_fact.id_prest.int_prest)
				paragraph3 = row_cells[2].paragraphs[0]
				run3 = paragraph3.add_run(dt_fr(obj_fact.dt_mand_moa_fact) or '-')
				paragraph4 = row_cells[3].paragraphs[0]
				run4 = paragraph4.add_run(obj_fact.num_bord_fact)
				paragraph5 = row_cells[4].paragraphs[0]
				run5 = paragraph5.add_run(obj_fact.num_mandat_fact)
				paragraph6 = row_cells[5].paragraphs[0]
				run6 = paragraph6.add_run(obt_mont(obj_fact.mont_ht_fact) or '-')
				paragraph7 = row_cells[6].paragraphs[0]
				run7 = paragraph7.add_run(obt_mont(obj_fact.mont_ttc_fact) or '-')

				# Mise en forme d'une ligne du tableau des factures
				for p in [paragraph1, paragraph2, paragraph3, paragraph4, paragraph5, paragraph6, paragraph7] :
					p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

				# Nettoyage de certaines variables (précaution)
				del paragraph1, paragraph2, paragraph3, paragraph4, paragraph5, paragraph6, paragraph7
				del run1, run2, run3, run4, run5, run6, run7

			# Préparation de la ligne "Total"
			footer_cells = table.add_row().cells
			paragraph1 = footer_cells[0].paragraphs[0]
			run1 = paragraph1.add_run('Total (en €)')
			run1.font.size = Pt(8)
			paragraph2 = footer_cells[5].paragraphs[0]
			run2 = paragraph2.add_run(obt_mont(mont_ht_fact_sum))
			paragraph3 = footer_cells[6].paragraphs[0]
			run3 = paragraph3.add_run(obt_mont(mont_ttc_fact_sum))

			# Mise en forme de la ligne "Total"
			for r in [run1, run2, run3] :
				r.bold = True
			for p in [paragraph1, paragraph2, paragraph3] :
				p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

			# Nettoyage de certaines variables (précaution)
			del paragraph1, paragraph2, paragraph3
			del run1, run2, run3

			# Mise en forme générale du tableau des factures
			for c in header_cells :
				c._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="F6F6F6"/>'.format(nsdecls('w'))))
			for c in footer_cells :
				c._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="F6F6F6"/>'.format(nsdecls('w'))))

			# Saut de ligne
			document.add_paragraph()

		# Déclaration du tableau des tampons
		table = document.add_table(rows = 0, cols = 2)
		row_cells = table.add_row().cells

		# Insertion du tampon de la trésorerie
		paragraph = row_cells[0].paragraphs[0]
		run = paragraph.add_run('À [COMMUNE], le')
		run.add_break()
		run = paragraph.add_run('Le trésorier de [COMMUNE],')

		# Insertion du tampon du syndicat
		paragraph = row_cells[1].paragraphs[0]
		run = paragraph.add_run('À [COMMUNE], le')
		run.add_break()
		run = paragraph.add_run('Le président du/de la/de l\'/des {},'.format(obj_org_moa))

		# Sauvegarde du document Word et téléchargement de celui-ci
		output = HttpResponse(content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
		output['Content-Disposition'] = 'attachment; filename={}.docx'.format(gen_cdc())
		document.save(output)

	return output

'''
Cette vue permet de traiter le formulaire de suppression d'un élément composant la fiche de vie d'un dossier.
rq : Objet requête
_fdv : Identifiant d'un événement
'''
@verif_acc
@csrf_exempt
def suppr_fdv(rq, _fdv) :

	# Imports
	from app.functions import ger_droits
	from app.functions import rempl_fich_log
	from app.models import TFicheVie
	from django.core.urlresolvers import reverse
	from django.http import HttpResponse
	from django.shortcuts import get_object_or_404
	import json

	output = HttpResponse()

	# Obtention d'une instance TFicheVie
	fdv = get_object_or_404(TFicheVie, pk = _fdv)

	if rq.method == 'POST' :

		# Vérification du droit d'écriture
		ger_droits(rq.user, fdv.id_doss, False)

		# Pointage vers une instance TDossier
		dos = fdv.id_doss

		# Récupération de l'identifiant de l'événement afin de le renseigner dans le fichier log
		id_fdv = fdv.pk

		# Suppression d'un événement
		fdv.delete()

		# Affichage d'un message de succès
		output = HttpResponse(
			json.dumps({ 'success' : {
				'message' : 'La fiche de vie du dossier {} a été mise à jour avec succès.'.format(dos),
				'redirect' : reverse('cons_doss', args = [dos.pk])
			}}),
			content_type = 'application/json'
		)

		# Définition de l'onglet actif après rechargement de la page
		rq.session['tab_doss'] = '#ong_fdv'

		# Remplissage du fichier log
		rempl_fich_log([rq.user.pk, rq.user, dos.pk, dos, 'D', 'Fiche de vie', id_fdv])

	return output

'''
Cette vue permet d'afficher la mise à jour d'un élément composant la fiche de vie
rq : Objet requête
_fdv : Identifiant d'un événement
'''
@verif_acc
def modif_fdv(request, pk) :

	output = HttpResponse()

	# Obtention d'une instance TFicheVie
	fdv = get_object_or_404(TFicheVie, pk=pk)

	# Vérification du droit d'écriture
	ger_droits(request.user, fdv.id_doss, False)

	if request.method == 'GET' :

		# J'instancie un objet "formulaire".
		form = GererFicheVie(instance = fdv, k_doss = fdv.id_doss)

		# Affichage du template
		output = render(request, './gestion_dossiers/modif_fdv.html', {
			'fdv' : fdv,
			'form' : form.get_form(rq=request),
			't_fm' : [init_fm('gerfdv', 'Mettre à jour un événement')],
			'title' : 'Mettre à jour un événement',
		})

	else :

		# Soumission du formulaire
		form = GererFicheVie(request.POST, request.FILES, instance = fdv, k_doss = fdv.id_doss)

		if form.is_valid() :

			# Mise à jour d'une instance TFicheVie
			fdv = form.save()

			# Affichage du message de succès
			output = HttpResponse(
				json.dumps({ 'success' : {
					'message' : 'La fiche de vie du dossier {} a été mise à jour avec succès.'.format(fdv.id_doss),
					'redirect' : reverse('cons_doss', args = [fdv.id_doss.pk])
				}}),
				content_type = 'application/json'
			)

			# Définition de l'onglet actif après rechargement de la page
			request.session['tab_doss'] = '#ong_fdv'

			# Remplissage du fichier log
			rempl_fich_log([
				request.user.pk,
				request.user,
				fdv.id_doss.pk,
				fdv.id_doss,
				'U',
				'Fiche de vie',
				fdv.pk
			])

		else :

			# Affichage des erreurs
			output = HttpResponse(json.dumps(form.errors), content_type = 'application/json')

	return output

'''
Cette vue permet de traiter le formulaire d'ajout d'un ordre de service.
request : Objet requête
pd_d = Identifiant d'une instance TPrestationsDossier
p_redirect : Paramètres de redirection
'''
@verif_acc
def ajout_os(request, pd_id, p_redirect) :

	# Imports
	from app.forms.gestion_dossiers import GererOrdreService
	from app.functions import ger_droits
	from app.functions import rempl_fich_log
	from app.models import TPrestationsDossier
	from django.core.urlresolvers import reverse
	from django.http import HttpResponse
	from django.shortcuts import get_object_or_404
	import json

	output = HttpResponse()

	if request.method == 'POST' :

		# Obtention d'une instance TPrestationsDossier
		pd = get_object_or_404(TPrestationsDossier, pk = pd_id)

		# Je vérifie le droit d'écriture.
		ger_droits(request.user, pd.id_doss, False)

		# Je soumets le formulaire.
		form = GererOrdreService(request.POST, k_pd = pd, prefix = 'GererOrdreService')

		if form.is_valid() :

			# Création d'une instance TOrdreService
			os = form.save()

			# Définition des paramètres de redirection
			if p_redirect == 'cons_prest' :
				prms = { 'url' : reverse('cons_prest', args = [pd.pk]), 'tab' : 'tab_prest', 'ong' : '#ong_os' }
			elif p_redirect == 'cons_doss' :
				prms = { 'url' : reverse('cons_doss', args = [os.id_doss.pk]), 'tab' : 'tab_doss', 'ong' : '#ong_prest' }
			else :
				prms = {}

			if prms :

				# J'affiche le message de succès.
				output = HttpResponse(
					json.dumps({ 'success' : {
						'message' : '''
						L'ordre de service a été ajouté avec succès à la prestation « {} » du dossier {}.
						'''.format(os.id_prest, os.id_doss),
						'redirect' : prms['url']
					}}),
					content_type = 'application/json'
				)

				# Je renseigne l'onglet actif après rechargement de la page.
				request.session[prms['tab']] = prms['ong']

				# Je complète le fichier log.
				rempl_fich_log([
					request.user.pk, request.user, os.id_doss.pk, os.id_doss, 'C', 'Ordre de service', os.pk
				])

		else :

			# J'affiche les erreurs.
			t_err = {}
			for k, v in form.errors.items() :
				t_err['GererOrdreService-{0}'.format(k)] = v
			output = HttpResponse(json.dumps(t_err), content_type = 'application/json')

	return output

'''
Cette vue permet de traiter le formulaire de suppression d'un ordre de service.
rq : Objet requête
os_id : Identifiant d'un ordre de service
'''
@verif_acc
@csrf_exempt
def suppr_os(rq, os_id) :

	# Imports
	from app.functions import ger_droits
	from app.functions import rempl_fich_log
	from app.models import TOrdreService
	from app.models import TPrestationsDossier
	from django.core.urlresolvers import reverse
	from django.http import HttpResponse
	from django.shortcuts import get_object_or_404
	import json

	output = HttpResponse()

	# Obtention d'une instance TOrdreService
	os = get_object_or_404(TOrdreService, pk = os_id)

	if rq.method == 'POST' :

		# Vérification du droit d'écriture
		ger_droits(rq.user, os.id_doss, False)

		# Pointage vers une instance TPrestationsDossier
		pd = TPrestationsDossier.objects.get(id_doss = os.id_doss, id_prest = os.id_prest)

		# Récupération de l'identifiant de l'ordre de service afin de le renseigner dans le fichier log
		id_os = os.pk

		# Suppression d'un ordre de service
		os.delete()

		# Affichage d'un message de succès
		output = HttpResponse(
			json.dumps({ 'success' : {
				'message' : 'L\'ordre de service a été supprimé avec succès.',
				'redirect' : reverse('cons_prest', args = [pd.pk])
			}}),
			content_type = 'application/json'
		)

		# Définition de l'onglet actif après rechargement de la page
		rq.session['tab_prest'] = '#ong_os'

		# Remplissage du fichier log
		rempl_fich_log([rq.user.pk, rq.user, pd.id_doss.pk, pd.id_doss, 'D', 'Ordre de service', id_os])

	return output

'''
Cette vue permet la mise à jour d'un ordre de service
rq : Objet requête
os_id : Identifiant d'un ordre de service
'''
@verif_acc
def modif_os(request, os_id) :

	# Imports
	from app.forms.gestion_dossiers import GererOrdreService
	from app.functions import ger_droits
	from app.functions import init_fm
	from app.functions import rempl_fich_log
	from app.models import TOrdreService
	from app.models import TPrestationsDossier
	from django.core.urlresolvers import reverse
	from django.http import HttpResponse
	from django.shortcuts import get_object_or_404
	from django.shortcuts import render
	import json

	output = HttpResponse()

	# Obtention d'une instance TOrdreService
	os = get_object_or_404(TOrdreService, pk = os_id)

	# Vérification du droit d'écriture
	ger_droits(request.user, os.id_doss, False)

	# Obtention d'une instance TPrestationsDossier
	pd = TPrestationsDossier.objects.get(id_doss = os.id_doss, id_prest = os.id_prest)

	if request.method == 'GET' :

		# J'instancie un objet "formulaire".
		form = GererOrdreService(instance = os)

		# Affichage du template
		output = render(request, './gestion_dossiers/modif_os.html', {
			'form' : form.get_form(rq = request),
			'pd' : pd,
			't_fm' : [init_fm('geros', 'Modifier un ordre de service')],
			'title' : 'Modifier un ordre de service',
		})

	else :

		# Soumission du formulaire
		form = GererOrdreService(request.POST, instance = os)

		if form.is_valid() :

			# Mise à jour d'une instance TOrdreService
			os = form.save()

			# Affichage du message de succès
			output = HttpResponse(
				json.dumps({ 'success' : {
					'message' : 'L\'ordre de service a été mis à jour avec succès.',
					'redirect' : reverse('cons_prest', args = [pd.pk])
				}}),
				content_type = 'application/json'
			)

			# Définition de l'onglet actif après rechargement de la page
			request.session['tab_prest'] = '#ong_os'

			# Remplissage du fichier log
			rempl_fich_log([
				request.user.pk,
				request.user,
				os.id_doss.pk,
				os.id_doss,
				'U',
				'Ordre de service',
				os.pk
			])

		else :

			# Affichage des erreurs
			output = HttpResponse(json.dumps(form.errors), content_type = 'application/json')

	return output